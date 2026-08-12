"""Phase 3 新增逻辑测试: REPEAT_BLOCK 块复制 + 章节裁剪。"""

from types import SimpleNamespace

from docx import Document

from apps.response_template.constants import BlockFillStatus
from apps.response_template.services.filler import OoxmlFiller

filler = OoxmlFiller()


def make_block(**kw):
    defaults = {
        "block_key": "R1",
        "title": "t",
        "anchor_text": "",
        "binding_config": {"repeat_count": 3},
        "fill_payload": {},
        "save": lambda **k: None,
    }
    defaults.update(kw)
    return SimpleNamespace(**defaults)


class TestRepeatBlock:
    def test_copy_block_three_times(self):
        doc = Document()
        doc.add_paragraph("二、项目成员简历")
        doc.add_paragraph("姓名：____")
        doc.add_paragraph("工作年限：____ 年")
        doc.add_paragraph("三、项目成员社保证明")
        doc.add_paragraph("附件8 其他")

        block = make_block(anchor_text="二、项目成员简历")
        warnings = []
        status = filler._repeat_block(doc, block, warnings)
        assert status == BlockFillStatus.FILLED

        texts = [p.text for p in doc.paragraphs]
        # 锚点标题 1 份 + 内容复制 3 份(每份 2 段)
        assert texts.count("姓名：____") == 3
        assert texts.count("工作年限：____ 年") == 3
        assert not warnings
        assert block.fill_payload["copied"] == 3

    def test_copy_respects_repeat_count(self):
        doc = Document()
        doc.add_paragraph("简历模板")
        doc.add_paragraph("姓名：____")
        doc.add_paragraph("二、其他内容")

        block = make_block(anchor_text="简历模板", binding_config={"repeat_count": 2})
        warnings = []
        status = filler._repeat_block(doc, block, warnings)
        assert status == BlockFillStatus.FILLED
        texts = [p.text for p in doc.paragraphs]
        assert texts.count("姓名：____") == 2

    def test_missing_anchor(self):
        doc = Document()
        doc.add_paragraph("其他内容")
        block = make_block(anchor_text="不存在的锚点")
        warnings = []
        status = filler._repeat_block(doc, block, warnings)
        assert status == BlockFillStatus.NEEDS_REVIEW
        assert warnings


class TestTrimToAnchor:
    def test_trim_removes_before_anchor(self):
        doc = Document()
        doc.add_paragraph("第一部分 前面内容")
        doc.add_paragraph("附件7：响应价格一览表")
        doc.add_paragraph("报价内容")

        ok = filler._trim_to_anchor(doc, "附件7")
        assert ok is True
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["附件7：响应价格一览表", "报价内容"]

    def test_trim_missing_anchor(self):
        doc = Document()
        doc.add_paragraph("只有一段")
        ok = filler._trim_to_anchor(doc, "不存在的锚点")
        assert ok is False
        assert len(doc.paragraphs) == 1
