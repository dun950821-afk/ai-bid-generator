"""填充引擎测试(内存 docx, 不碰 DB)。"""

from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from apps.response_template.constants import BlockFillStatus
from apps.response_template.services.filler import OoxmlFiller

filler = OoxmlFiller()


def make_doc(paragraphs=None, tables=None):
    """构造内存 docx。"""
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    for rows in tables or []:
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
    return doc


def make_block(**kw):
    defaults = {
        "block_key": "T1",
        "title": "t",
        "anchor_text": "",
        "binding_config": {},
        "source_config": {},
        "block_type": "AUTO_FIELD",
    }
    defaults.update(kw)
    return SimpleNamespace(**defaults)


class TestPlaceholderReplace:
    def test_underline(self):
        assert filler._replace_first_placeholder("地址：____", "北京") == "地址：北京"
        assert filler._replace_first_placeholder("____年", "2026") == "2026年"

    def test_year_month_day(self):
        text = "日    期：      年     月     日"
        out = filler._replace_first_placeholder(text, "2026年8月12日")
        assert out == "日    期：2026年8月12日"

    def test_year_month_day_no_double_fill(self):
        """已填充的日期不再被识别为占位(回归: 20262026年08月12日)。"""
        text = "日    期：2026年08月12日"
        out = filler._replace_first_placeholder(text, "2026年08月13日", hint_text="日期")
        assert out == text

    def test_paren_empty(self):
        assert filler._replace_first_placeholder("（  ）", "010-1234") == "010-1234"

    def test_paren_hint_requires_hint(self):
        # 提示括号需 hint_text 匹配才替换
        assert filler._replace_first_placeholder("（电话）", "010-1234", hint_text="电话") == "010-1234"
        assert filler._replace_first_placeholder("（电话）", "x") == "（电话）"

    def test_paren_hint_protects_body_text(self):
        """（法人公章）（大写）等正文括号不替换, 只填真正的空位。"""
        text = "响应人（法人公章）："
        assert filler._replace_first_placeholder(text, "X公司", hint_text="响应人（法人公章）") == text
        out = filler._replace_first_placeholder("人民币(大写)        元整", "10000", hint_text="人民币(大写)")
        assert out == "人民币(大写)10000元整"

    def test_space_run(self):
        text = "根据贵方            项目采购文件"
        out = filler._replace_first_placeholder(text, "渗透测试1", hint_text="根据贵方")
        assert out == "根据贵方渗透测试1项目采购文件"

    def test_hint_paren_swallows_leading_spaces(self):
        text = "响应人                  （响应人名称）提供相关文件"
        out = filler._replace_first_placeholder(text, "X公司", hint_text="响应人 （响应人名称）")
        assert out == "响应人X公司提供相关文件"

    def test_no_placeholder_unchanged(self):
        text = "普通文本"
        assert filler._replace_first_placeholder(text, "x") == text


class TestNormalize:
    def test_fullwidth_paren(self):
        assert filler._normalize("响应人 （名称）") == "响应人(名称)"
        assert filler._normalize("地址 ： xx") == "地址：xx"


class TestFindParagraph:
    def test_prefers_placeholder_paragraph(self):
        doc = make_doc(paragraphs=[
            "与本响应有关的一切正式信函使用以下地址：",  # 先出现, 无空位
            "地址：              （地址）",             # 有空位
        ])
        para = filler._find_paragraph(doc, "地址")
        assert "（地址）" in para.text

    def test_paragraph_in_table_cell(self):
        doc = make_doc(tables=[[["响应人名称", "____"], ["经营范围", "____"]]])
        para = filler._find_paragraph(doc, "经营范围")
        assert para is not None
        assert "经营范围" in para.text


class TestTextFill:
    def test_fill_underline(self):
        doc = make_doc(paragraphs=["电话：____"])
        block = make_block(anchor_text="电话")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "010-12345678", warnings)
        assert status == BlockFillStatus.FILLED
        assert "010-12345678" in doc.paragraphs[0].text
        assert not warnings

    def test_fill_year_month_day(self):
        doc = make_doc(paragraphs=["日    期：      年     月     日"])
        block = make_block(anchor_text="日期")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "2026年8月12日", warnings)
        assert status == BlockFillStatus.FILLED
        assert "2026年8月12日" in doc.paragraphs[0].text

    def test_missing_anchor(self):
        doc = make_doc(paragraphs=["其他内容"])
        block = make_block(anchor_text="不存在的锚点")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "v", warnings)
        assert status == BlockFillStatus.NEEDS_REVIEW
        assert warnings


class TestTableCellFill:
    def test_fill_row_after_label(self):
        doc = make_doc(tables=[[["响应人名称", "____"], ["经营范围", "____"]]])
        block = make_block(anchor_text="响应人名称")
        warnings = []
        status = filler._fill_table_by_label(doc, block, "某某科技有限公司", warnings)
        assert status == BlockFillStatus.FILLED
        assert "某某科技有限公司" in doc.tables[0].cell(0, 1).text

    def test_skips_tblpr(self):
        """含 tblPrEx 的复杂行不应报错(回归: CT_TblPrEx p_lst)。"""
        doc = make_doc(tables=[[["响应人名称", "____"], ["经营范围", "____"]]])
        block = make_block(anchor_text="响应人名称")
        warnings = []
        status = filler._fill_table_by_label(doc, block, "v", warnings)
        assert status == BlockFillStatus.FILLED


class TestRowClone:
    def test_clone_row(self):
        doc = make_doc(tables=[[["项目阶段", "工作项"], ["启动", ""], ["实施", ""]]])
        table = doc.tables[0]
        filler._clone_table_row(table, table.rows[-1])
        assert len(table.rows) == 4
        # 新行文本为空
        assert table.rows[-1].cells[0].text == ""


class TestTrim:
    def test_trim_before_section(self):
        """主文件只保留"响应文件格式"章节起的内容。"""
        doc = make_doc(paragraphs=[
            "第一部分 采购公告",
            "项目需求说明",
            "第四部分 响应文件格式",
            "附件1：响应文件",
        ])
        assert filler._trim_before_section(doc) is True
        texts = [p.text for p in doc.paragraphs]
        assert texts[0] == "第四部分 响应文件格式"
        assert "项目需求说明" not in texts

    def test_trim_before_section_not_found(self):
        doc = make_doc(paragraphs=["普通文档内容"])
        assert filler._trim_before_section(doc) is False

    def test_remove_attachments(self):
        doc = make_doc(paragraphs=[
            "第四部分 响应文件格式",
            "附件1：响应文件",
            "附件1 内容",
            "附件7：报价表",
            "报价内容",
            "附件8：其他文件",
            "其他内容",
        ])
        removed = filler._remove_attachments(doc, ["7"])
        assert removed == 2
        texts = [p.text for p in doc.paragraphs]
        assert "附件7：报价表" not in texts
        assert "报价内容" not in texts
        assert "附件1 内容" in texts
        assert "附件8：其他文件" in texts

    def test_keep_only_attachments(self):
        doc = make_doc(paragraphs=[
            "第四部分 响应文件格式",
            "附件7：报价表",
            "报价内容",
            "附件8：其他文件",
            "其他内容",
        ])
        removed = filler._keep_only_attachments(doc, ["7"])
        assert removed == 2
        texts = [p.text for p in doc.paragraphs]
        assert texts[0] == "第四部分 响应文件格式"  # 首个附件标题前的内容保留
        assert "报价内容" in texts
        assert "附件8：其他文件" not in texts


class TestNoAppendOnMissingPlaceholder:
    def test_no_placeholder_marks_review(self):
        """锚点段落无空位时不再追加到句尾(回归: 公司名错位)。"""
        doc = make_doc(paragraphs=["响应人完全接受采购文件中的内容。"])
        block = make_block(anchor_text="响应人完全接受")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "X公司", warnings)
        assert status == BlockFillStatus.NEEDS_REVIEW
        assert "X公司" not in doc.paragraphs[0].text
        assert warnings
