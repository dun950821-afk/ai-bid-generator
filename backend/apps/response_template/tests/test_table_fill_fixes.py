"""表格填充修复测试: 整行表头匹配 + 行内空 cell 填充 + AI 表格兜底。"""

from types import SimpleNamespace

from docx import Document

from apps.response_template.constants import BlockFillStatus
from apps.response_template.services.filler import OoxmlFiller

filler = OoxmlFiller()


def make_block(**kw):
    defaults = {
        "block_key": "T1",
        "title": "t",
        "anchor_text": "",
        "binding_config": {},
        "source_config": {},
        "fill_payload": {},
        "save": lambda **k: None,
    }
    defaults.update(kw)
    return SimpleNamespace(**defaults)


class TestTableHeaderMatch:
    def test_find_table_by_full_row(self):
        """表头跨多 cell 时整行拼接匹配。"""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        # 表头: 三个 cell 分开
        table.rows[0].cells[0].text = "采购文件章节号"
        table.rows[0].cells[1].text = "要求描述"
        table.rows[0].cells[2].text = "响应人响应情况"
        table.rows[1].cells[0].text = ""

        t = filler._find_table(doc, "采购文件章节号 要求描述 响应人响应情况")
        assert t is not None and t._tbl is table._tbl

    def test_find_table_by_single_cell(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "项目阶段"
        t = filler._find_table(doc, "项目阶段")
        assert t is not None and t._tbl is table._tbl


class TestRowEmptyCell:
    def test_fill_row_empty_cell(self):
        """label cell 无空位时填同行其他空 cell, 不覆盖 label。"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "响应人名称"
        table.rows[0].cells[1].text = ""
        block = make_block(anchor_text="响应人名称")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "北京国舜科技股份有限公司", warnings)
        assert status == BlockFillStatus.FILLED
        assert table.rows[0].cells[0].text == "响应人名称"  # label 未被覆盖
        assert "北京国舜" in table.rows[0].cells[1].text  # 值填到空 cell

    def test_no_empty_cell_keeps_label(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "只有 label"
        block = make_block(anchor_text="只有 label")
        warnings = []
        status = filler._fill_text_placeholder(doc, block, "值", warnings)
        # 表格 cell 无空位 → 不追加(避免污染 label), 标记人工
        assert status == BlockFillStatus.NEEDS_REVIEW
        assert table.rows[0].cells[0].text == "只有 label"


class TestAiTableFallback:
    def test_fill_first_data_row(self):
        """AI_GENERATE 兜底: 表头下第一数据行填充。"""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "项目阶段"
        table.rows[0].cells[1].text = "工作项"
        table.rows[0].cells[2].text = "主要交付物"
        table.rows[1].cells[0].text = ""

        block = make_block(anchor_text="项目阶段 工作项 主要交付物")
        warnings = []
        status = filler._fill_first_data_row(doc, block, "项目启动：实施准备", warnings)
        assert status == BlockFillStatus.FILLED
        assert table.rows[1].cells[0].text == "项目启动：实施准备"

    def test_fill_first_data_row_adds_row(self):
        """只有表头行时复制一行再填充。"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "项目阶段"
        table.rows[0].cells[1].text = "交付物"

        block = make_block(anchor_text="项目阶段 交付物")
        warnings = []
        status = filler._fill_first_data_row(doc, block, "启动阶段内容", warnings)
        assert status == BlockFillStatus.FILLED
        assert len(table.rows) == 2
        assert table.rows[1].cells[0].text == "启动阶段内容"
