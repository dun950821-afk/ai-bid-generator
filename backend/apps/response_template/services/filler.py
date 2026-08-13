# -*- coding: utf-8 -*-
"""OoxmlFiller v1: 响应模板原位填充引擎。

原则: 修改现有 XML, 不重建 Word。在招标方原始 docx 上按"文本锚点"定位填充,
保留表格边框/列宽/字体/底纹/合并单元格等全部格式。

填充策略(按 block_type):
- FIXED            → 跳过
- AUTO_FIELD       → 文本空位替换(段落/表格cell, 下划线/括号占位)
- DATA_TABLE       → 按行标签定位表格行, 填充空 cell
- MANUAL / PRICE   → 保留原文, 标记 needs_review(生成后由用户填写)
- AI_GENERATE      → LLM 生成内容后替换空位
- AI_RESPONSE      → 从招标条款生成逐条应答, 填充应答表格
- REPEAT_TABLE     → deepcopy 模板数据行 N 份
- MATERIAL_SLOT    → 从材料包取图插入

定位方式 v1(已知限制): 文本锚点 + 归一化匹配(去空白)。
v2 将升级为 Content Control 编译定位(复用 outline template_compiler 协议)。

数据源:
- 企业字段: CompanyProfile(is_default) + Project
- 材料: BidMaterialPackage.get_material_by_usage_key
- AI 内容: LLMService(DeepSeek), 与 generation 共用同一模型配置
"""

import io
import json
import logging
import re
from typing import List, Optional, Tuple

from django.core.files.base import ContentFile

from apps.common.services.storage import StorageService
from apps.response_template.constants import BlockFillStatus, BlockType

logger = logging.getLogger(__name__)

# 空位模式
UNDERLINE_RE = re.compile(r"_{2,}")
# 空白/下划线括号占位: （ ）（＿＿）等(内容仅空白或下划线)
PAREN_EMPTY_RE = re.compile(r"[（(][_\s]{0,15}[）)]")
# 提示括号占位: （邮编）（响应人名称）等(内容 ≤12 字, 仅在与锚点匹配时才替换,
# 避免误伤"（法人公章）""（或授权代表人）"等正文括号)
PAREN_HINT_RE = re.compile(r"[（(]\s*[^）)]{1,12}[）)]")
# 连续空格占位(6+ 半角空格或 2+ 全角空格): "根据贵方      项目采购文件"
SPACE_RUN_RE = re.compile(r" {6,}|　{2,}")
# "年  月  日" 空格占位(落款日期常见形态);
# 负向后查 (?<!\d): 已填充的"2026年08月12日"不再被误判为占位(防止重复填充)
YEAR_MONTH_DAY_RE = re.compile(r"(?<!\d)\s*年\s*\d{0,2}\s*月\s*\d{0,2}\s*日")

# 图片扩展名
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}

# 应答状态枚举(风险闸门)
RESPONSE_VALUES = ["完全响应", "部分响应", "偏离", "待确认"]


class FillWarning:
    """填充警告(不阻断生成)。"""

    def __init__(self, block_key: str, message: str):
        self.block_key = block_key
        self.message = message

    def to_dict(self) -> dict:
        return {"block_key": self.block_key, "message": self.message}


class OoxmlFiller:
    """响应模板填充引擎。"""

    def __init__(self):
        self.storage = StorageService()

    # ------------------------------------------------------------------
    # 对外入口
    # ------------------------------------------------------------------
    def fill(
        self, template, blocks: List, trim_anchor: Optional[str] = None,
        trim_section: bool = False, exclude_attachments: Optional[List[str]] = None,
        keep_only_attachments: Optional[List[str]] = None,
    ) -> Tuple[ContentFile, List[dict], List]:
        """填充原始 docx。

        Args:
            template: 响应模板
            blocks: 待填充的块
            trim_anchor: 若指定, 生成后删除该锚点之前的所有内容
                        (用于单独密封文档按章节裁剪)
            trim_section: 若 True, 删除"第X部分 响应文件格式"章节之前的所有内容
                        (主响应文件只保留响应格式部分, 不含招标文件正文)
            exclude_attachments: 要从产物中删除的附件编号列表
                        (单独密封附件不进主文件)
            keep_only_attachments: 只保留指定编号的附件, 其余附件删除
                        (密封文档只含密封附件)

        Returns:
            (ContentFile, warnings, filled_blocks)
        """
        from docx import Document

        # 优先使用 compiled 模板(带 Content Control 标记), 否则原始文件
        object_key = getattr(template, "compiled_file_key", "") or template.source_file.object_key
        raw = self.storage.get_object(object_key)
        doc = Document(io.BytesIO(raw))

        company, project = self._load_data_sources(template)
        material_package = self._load_material_package(template)
        warnings: List[FillWarning] = []
        filled: List = []

        # 有序填充
        ordered = list(blocks)
        for block in ordered:
            try:
                status = self._fill_block(doc, block, company, project, material_package, warnings)
                block.fill_status = status
                block.save(update_fields=["fill_status", "updated_at"])
                if status != BlockFillStatus.SKIPPED:
                    filled.append(block)
            except Exception as exc:
                logger.warning("block fill failed: key=%s err=%s", block.block_key, exc)
                block.fill_status = BlockFillStatus.NEEDS_REVIEW
                block.save(update_fields=["fill_status", "updated_at"])
                warnings.append(FillWarning(block.block_key, f"填充失败: {exc}"))

        # 裁剪主文档: 只保留"响应文件格式"章节起的内容
        if trim_section:
            if not self._trim_before_section(doc):
                warnings.append(FillWarning("TRIM", "未找到'响应文件格式'章节, 保留完整文档"))
        # 按章节裁剪(单独密封文档)
        if trim_anchor:
            trimmed = self._trim_to_anchor(doc, trim_anchor)
            if not trimmed:
                warnings.append(FillWarning("TRIM", f"未找到裁剪锚点: {trim_anchor}, 保留完整文档"))
        # 删除单独密封附件(不进主文件)
        if exclude_attachments:
            self._remove_attachments(doc, exclude_attachments)
        # 只保留密封附件(密封文档用)
        if keep_only_attachments:
            self._keep_only_attachments(doc, keep_only_attachments)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"{template.name or '响应文件'}.docx"
        content_file = ContentFile(buffer.read(), name=filename)
        return content_file, [w.to_dict() for w in warnings], filled

    # 章节标题模式: "第四部分 响应文件格式"
    SECTION_TITLE_RE = re.compile(r"^第[一二三四五六七八九十百\d]+部分")

    def _trim_before_section(self, doc) -> bool:
        """删除"响应文件格式"章节之前的所有 body 元素(保留 sectPr)。

        主响应文件只需要响应格式部分, 招标文件正文(申明/需求/采购说明)不应出现。
        """
        body = doc.element.body
        target = None
        for el in body.iterchildren():
            text = self._normalize("".join(el.itertext()))
            if "响应文件格式" in text and self.SECTION_TITLE_RE.match(text):
                target = el
                break
        if target is None:
            # 兜底: 任意短标题含"响应文件格式"
            for el in body.iterchildren():
                text = self._normalize("".join(el.itertext()))
                if "响应文件格式" in text and len(text) <= 30:
                    target = el
                    break
        if target is None:
            return False
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for el in list(body):
            if el is target:
                break
            if el.tag == f"{w_ns}sectPr":
                continue
            body.remove(el)
        return True

    def _section_scoped_children(self, doc):
        """返回 (body 子元素列表, 响应格式章节之后的起始下标)。

        招标文件正文前的指引/目录可能也列有"附件N"标题, 附件删除操作
        必须限定在"响应文件格式"章节之后, 避免误删正文。
        """
        body = doc.element.body
        children = list(body.iterchildren())
        start = 0
        for i, el in enumerate(children):
            text = self._normalize("".join(el.itertext()))
            if "响应文件格式" in text and (
                self.SECTION_TITLE_RE.match(text) or len(text) <= 30
            ):
                start = i + 1
                break
        return children, start

    def _remove_attachments(self, doc, attachment_nos) -> int:
        """删除指定编号的附件(从"附件N"标题到下一附件标题前)。返回删除的元素数。"""
        targets = {str(n) for n in attachment_nos if n}
        if not targets:
            return 0
        body = doc.element.body
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        children, start = self._section_scoped_children(doc)
        removing = False
        removed = 0
        for el in children[start:]:
            if el.tag == f"{w_ns}sectPr":
                continue
            text = self._normalize("".join(el.itertext()))
            m = re.match(r"附件(\d+)", text)
            if m:
                removing = m.group(1) in targets
            if removing:
                body.remove(el)
                removed += 1
        return removed

    def _keep_only_attachments(self, doc, keep_nos) -> int:
        """只保留指定编号的附件, 删除其余附件(首个附件标题前的内容保留)。"""
        keeps = {str(n) for n in keep_nos if n}
        body = doc.element.body
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        children, start = self._section_scoped_children(doc)
        removing = False
        removed = 0
        for el in children[start:]:
            if el.tag == f"{w_ns}sectPr":
                continue
            text = self._normalize("".join(el.itertext()))
            m = re.match(r"附件(\d+)", text)
            if m:
                removing = m.group(1) not in keeps
            if removing:
                body.remove(el)
                removed += 1
        return removed

    def _trim_to_anchor(self, doc, anchor_text: str) -> bool:
        """删除 anchor 之前的所有 body 元素(保留 sectPr), 实现章节裁剪。

        从"第X部分 ...响应文件格式"章节之后定位 anchor,
        避免命中目录/指引里同名的标题。
        """
        anchor_norm = self._normalize(anchor_text)
        if not anchor_norm:
            return False
        body = doc.element.body
        anchor_el = None
        section_el = None
        started = False
        for el in body.iterchildren():
            text = self._normalize("".join(el.itertext()))
            if not started:
                if "响应文件格式" in text and re.match(r"^第.+部分", text):
                    started = True
                    section_el = el  # 保留章节标题
                continue
            if anchor_norm in text:
                anchor_el = el
                break
        if anchor_el is None:
            return False
        to_remove = []
        for el in body:
            if el is anchor_el:
                break
            if el is section_el:
                continue  # 保留章节标题
            if el.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
                continue
            to_remove.append(el)
        for el in to_remove:
            body.remove(el)
        return True

    # ------------------------------------------------------------------
    # 数据源
    # ------------------------------------------------------------------
    def _load_data_sources(self, template):
        from apps.enterprise.models import CompanyProfile

        company = CompanyProfile.objects.filter(is_default=True).first()
        project = template.project
        return company, project

    def _load_material_package(self, template):
        """取材料包: 优先默认企业的第一个材料包。"""
        try:
            from apps.enterprise.models import CompanyProfile

            company = CompanyProfile.objects.filter(is_default=True).first()
            if company:
                return company.material_packages.first()
        except Exception:
            logger.exception("material package load failed")
        return None

    # ------------------------------------------------------------------
    # 块分发
    # ------------------------------------------------------------------
    def _fill_block(self, doc, block, company, project, material_package, warnings) -> str:
        btype = block.block_type
        if btype == BlockType.FIXED:
            return BlockFillStatus.SKIPPED

        if btype == BlockType.AUTO_FIELD:
            value = self._resolve_field_value(block, company, project, warnings)
            if value is None:
                # 无数据源: 清除（邮编）等提示占位, 不留残渣, 标记人工补填
                self._clear_hint_placeholders(doc, block)
                return BlockFillStatus.NEEDS_REVIEW
            return self._fill_text_placeholder(doc, block, value, warnings)

        if btype == BlockType.DATA_TABLE:
            value = self._resolve_field_value(block, company, project, warnings)
            if value is None:
                self._clear_hint_placeholders(doc, block)
                return BlockFillStatus.NEEDS_REVIEW
            return self._fill_table_by_label(doc, block, value, warnings)

        if btype == BlockType.AI_GENERATE:
            text = self._generate_text(block, project)
            if not text:
                warnings.append(FillWarning(block.block_key, "AI 生成内容为空"))
                return BlockFillStatus.NEEDS_REVIEW
            # 保存生成内容快照(填充失败时前端可复制人工粘贴)
            block.fill_payload = {**(block.fill_payload or {}), "generated_text": text}
            block.save(update_fields=["fill_payload", "updated_at"])
            status = self._fill_text_placeholder(doc, block, text, warnings)
            if status == BlockFillStatus.FILLED:
                return status
            # 兜底1: 表头跨多 cell 场景(如"项目阶段|工作项|主要交付物"),
            # 定位表格后填充第一数据行
            status = self._fill_first_data_row(doc, block, text, warnings)
            if status == BlockFillStatus.FILLED:
                return status
            # 兜底2: 锚点为普通段落时, 在其后插入新段落承载生成内容
            para = self._locate_paragraph(doc, block)
            if para is not None and not self._paragraph_in_cell(para):
                self._insert_paragraph_after(para, text)
                return BlockFillStatus.FILLED
            warnings.append(FillWarning(block.block_key, "AI 内容已生成但未找到填充位置, 请在块详情中复制"))
            return BlockFillStatus.NEEDS_REVIEW

        if btype == BlockType.AI_RESPONSE:
            return self._fill_response_table(doc, block, project, warnings)

        if btype == BlockType.REPEAT_TABLE:
            return self._repeat_table_rows(doc, block, warnings)

        if btype == BlockType.REPEAT_BLOCK:
            return self._repeat_block(doc, block, warnings)

        if btype == BlockType.MATERIAL_SLOT:
            return self._insert_material(doc, block, material_package, warnings)

        if btype in (BlockType.MANUAL, BlockType.PRICE):
            # PRICE 块: 前端已录入报价(fill_payload.price)则填充, 否则保留人工
            if btype == BlockType.PRICE:
                price = (block.fill_payload or {}).get("price")
                if price not in (None, ""):
                    return self._fill_text_placeholder(doc, block, str(price), warnings)
            warnings.append(FillWarning(block.block_key, f"【人工填写】{block.title}"))
            return BlockFillStatus.NEEDS_REVIEW

        return BlockFillStatus.NEEDS_REVIEW

    # ------------------------------------------------------------------
    # 企业字段解析
    # ------------------------------------------------------------------
    def _resolve_field_value(self, block, company, project, warnings) -> Optional[str]:
        """按 binding_config.field 解析值。返回 None 表示无数据源。"""
        binding = block.binding_config or {}
        field = binding.get("field", "")
        if not field:
            warnings.append(FillWarning(block.block_key, f"未绑定数据源: {block.title}"))
            return None

        if field.startswith("company."):
            if not company:
                warnings.append(FillWarning(block.block_key, f"企业资料缺失: {block.title}"))
                return None
            attr = field.split(".", 1)[1]
            value = getattr(company, attr, None)
            if hasattr(value, "strftime"):
                value = value.strftime("%Y年%m月%d日")
            if value is None or value == "":
                warnings.append(FillWarning(block.block_key, f"企业字段为空: {field}"))
                return None
            return str(value)

        if field.startswith("project."):
            attr = field.split(".", 1)[1]
            value = getattr(project, attr, None) if project else None
            if value in (None, ""):
                # 项目字段缺失(如 bid_date) → 落款日期用今天
                if attr == "bid_date":
                    from django.utils import timezone

                    return timezone.now().strftime("%Y年%m月%d日")
                warnings.append(FillWarning(block.block_key, f"项目字段为空: {field}"))
                return None
            return str(value)

        warnings.append(FillWarning(block.block_key, f"未知绑定: {field}"))
        return None

    # ------------------------------------------------------------------
    # 文本空位替换(AUTO_FIELD / AI_GENERATE / PRICE)
    # ------------------------------------------------------------------
    def _fill_text_placeholder(self, doc, block, value: str, warnings) -> str:
        """找到含锚点文本的段落, 替换其中的空位(下划线/空格/括号占位)。"""
        anchor = block.anchor_text.strip()
        if not anchor:
            warnings.append(FillWarning(block.block_key, "缺少定位锚点"))
            return BlockFillStatus.NEEDS_REVIEW

        para = self._locate_paragraph(doc, block)
        if para is None:
            warnings.append(FillWarning(block.block_key, f"未找到锚点段落: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        text = para.text
        new_text = self._replace_first_placeholder(text, value, hint_text=anchor or block.title)
        if new_text != text:
            self._set_paragraph_text(para, new_text)
            return BlockFillStatus.FILLED

        # 表格 label cell 场景: 该行其他 cell 有空位 → 填空 cell(不覆盖 label)
        row_target = self._find_row_empty_cell(para)
        if row_target is not None:
            self._fill_tc_text(row_target, value)
            return BlockFillStatus.FILLED

        # 表格 cell 内无空位 → 不追加(避免污染 label), 标记人工
        if self._paragraph_in_cell(para):
            warnings.append(FillWarning(block.block_key, f"表格行无空位可填: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        # 无空位 → 不追加到句尾(避免错位污染), 标记人工补填
        warnings.append(FillWarning(block.block_key, f"锚点段落无空位可填: {anchor}"))
        return BlockFillStatus.NEEDS_REVIEW

    def _clear_hint_placeholders(self, doc, block) -> None:
        """字段无数据源时, 清除锚点段落中与锚点相关的提示括号占位(（邮编）等)。

        只清除内容出现在锚点文本中的括号, 不动"（法人公章）"等正文括号。
        """
        anchor = (block.anchor_text or "").strip() or (block.title or "")
        if not anchor:
            return
        para = self._locate_paragraph(doc, block)
        if para is None:
            return
        text = para.text
        changed = text
        for m in reversed(list(PAREN_HINT_RE.finditer(text))):
            if self._is_hint_paren_fillable(m.group(0), anchor):
                changed = changed[:m.start()] + changed[m.end():]
        if changed != text:
            self._set_paragraph_text(para, changed)

    def _insert_paragraph_after(self, para, text: str) -> None:
        """在锚点段落后插入新段落(复制段落格式), 承载 AI 生成内容。"""
        from copy import deepcopy

        new_p_el = deepcopy(para._p)
        self._set_lxml_paragraph_text(new_p_el, text)
        para._p.addnext(new_p_el)

    @staticmethod
    def _paragraph_in_cell(para) -> bool:
        p_el = para._p
        tc = p_el.getparent()
        while tc is not None and tc.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc":
            tc = tc.getparent()
        return tc is not None

    def _find_row_empty_cell(self, para):
        """若段落属于表格 cell, 返回该行**当前 cell 之后**第一个空 cell 的 lxml tc 元素。

        只向后查找: 避免多字段行(响应人名称|空|详细地址|空)中
        把值填到 label 之前的空 cell。
        """
        p_el = para._p
        tc = p_el.getparent()
        while tc is not None and tc.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc":
            tc = tc.getparent()
        if tc is None:
            return None
        el = tc.getnext()
        while el is not None:
            if el.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc":
                el = el.getnext()
                continue
            cell_text = "".join(el.itertext()).strip()
            if not cell_text or UNDERLINE_RE.search(cell_text):
                return el
            el = el.getnext()
        return None

    @staticmethod
    def _fill_tc_text(tc, text: str) -> None:
        """向表格 cell 的第一个段落写入文本。"""
        ps = tc.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        if ps:
            OoxmlFiller._set_lxml_paragraph_text(ps[0], text)

    # 提示括号可作为占位的内容关键词: （响应人名称）（邮编）等;
    # 不在表中的括号视为正文标注(如（大写）（法人公章）), 不替换
    PAREN_HINT_KEYWORDS = (
        "名称", "地址", "电话", "邮编", "传真", "邮箱", "姓名",
        "日期", "金额", "数量", "单价", "总价", "报价", "盖章", "签字",
    )

    def _is_hint_paren_fillable(self, paren_text: str, hint_text: str) -> bool:
        """判断提示括号是否为可填充占位: 内容与锚点相关, 且像"待填提示"。"""
        content_norm = self._normalize(paren_text).strip("()")
        if not content_norm:
            return False
        hint_norm = self._normalize(hint_text or "")
        if content_norm not in hint_norm:
            return False
        if hint_norm == self._normalize(paren_text):
            return True  # 锚点本身就是这个括号(如"（响应人名称）")
        return any(k in content_norm for k in self.PAREN_HINT_KEYWORDS)

    def _replace_first_placeholder(self, text: str, value: str, hint_text: str = None) -> str:
        """依次替换首个空位: 下划线 → 年月日空格 → 空括号 → 提示括号 → 连续空格。

        hint_text: 锚点/字段名, 仅当提示括号内容与其相关且像待填提示时才替换,
        防止误伤"（法人公章）""（大写）"等正文括号; 替换时一并吞掉括号前的空格占位。
        """
        if UNDERLINE_RE.search(text):
            return UNDERLINE_RE.sub(value, text, count=1)
        if YEAR_MONTH_DAY_RE.search(text):
            # 替换 "年 月 日" 前的空格部分为值
            return YEAR_MONTH_DAY_RE.sub(value, text, count=1)
        if PAREN_EMPTY_RE.search(text):
            return PAREN_EMPTY_RE.sub(value, text, count=1)
        if hint_text:
            m = PAREN_HINT_RE.search(text)
            if m and self._is_hint_paren_fillable(m.group(0), hint_text):
                # 连同括号前的空格占位一起替换
                m2 = re.search(r"[ 　]{0,20}" + re.escape(m.group(0)), text)
                if m2:
                    return text[:m2.start()] + value + text[m2.end():]
                return text[:m.start()] + value + text[m.end():]
        if SPACE_RUN_RE.search(text):
            return SPACE_RUN_RE.sub(value, text, count=1)
        return text

    def _set_paragraph_text(self, para, text: str) -> None:
        """整段重写文本(保留段落样式, run 级格式会归一)。"""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

    # ------------------------------------------------------------------
    # 表格按行标签填充(DATA_TABLE)
    # ------------------------------------------------------------------
    def _fill_table_by_label(self, doc, block, value: str, warnings) -> str:
        """定位含锚点文本的表格行, 填充该行第一个空 cell。"""
        anchor = block.anchor_text.strip()
        attachment_no = (block.source_config or {}).get("attachment_no")
        cell = self._find_label_cell(doc, anchor, attachment_no=attachment_no)
        if cell is None:
            warnings.append(FillWarning(block.block_key, f"未找到表格标签: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        # 该行中第一个空 cell(跳过标签 cell, 仅取 w:tc 元素)
        row = cell._tc.getparent()
        target = None
        for tc in row:
            if tc.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc":
                continue  # 跳过 trPr/tblPrEx 等非单元格元素
            if tc is cell._tc:
                continue
            cell_text = "".join(tc.itertext()).strip()
            if not cell_text or UNDERLINE_RE.search(cell_text):
                target = tc
                break
        if target is None:
            # 无空 cell → 追加到标签 cell 后面
            warnings.append(FillWarning(block.block_key, f"标签行无空 cell: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        from docx.table import _Cell

        target_cell = _Cell(target, cell._parent)
        self._set_cell_text(target_cell, value)
        return BlockFillStatus.FILLED

    def _set_cell_text(self, cell, text: str) -> None:
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        self._set_paragraph_text(para, text)

    # ------------------------------------------------------------------
    # AI 内容生成(AI_GENERATE / AI_RESPONSE)
    # ------------------------------------------------------------------
    def _generate_text(self, block, project) -> str:
        """AI 生成正文内容。上下文: 项目需求条款 + 块标题。"""
        from apps.generation.constants import ModelType
        from apps.generation.models import ModelConfig
        from apps.generation.services.llm_service import LLMService

        context = self._build_ai_context(block, project, limit=2000)
        system = "你是资深投标文件撰写专家。根据招标文件要求和项目信息, 撰写投标响应内容。要求专业、具体、可落地, 不要空话。直接输出纯文本正文内容, 禁止使用 #、*、- 等 markdown 标记, 不要任何解释。"
        user = f"招标要求与项目信息:\n{context}\n\n需要撰写的内容: {block.title}"
        resp = LLMService().chat(
            ModelConfig.objects.get(model_type=ModelType.CHAT, is_default=True, is_active=True),
            system,
            user,
        )
        return resp.text.strip() or ""

    def _fill_first_data_row(self, doc, block, text: str, warnings) -> str:
        """AI 生成内容兜底: 定位表格, 写入表头下第一数据行(行不足则复制表头行)。

        仅当锚点是表头行文本时生效, 避免普通行标签(如"主要股东")误触发表格填充。
        """
        table = self._find_table(doc, block.anchor_text)
        if table is None:
            warnings.append(FillWarning(block.block_key, f"未找到目标表格: {block.anchor_text}"))
            return BlockFillStatus.NEEDS_REVIEW
        header_idx, header_text = self._detect_header_row(table)
        if header_idx is None:
            return BlockFillStatus.NEEDS_REVIEW
        anchor_norm = self._normalize(block.anchor_text)
        header_norm = self._normalize(header_text)
        if not (anchor_norm and header_norm and (anchor_norm in header_norm or header_norm in anchor_norm)):
            # 锚点与表头不匹配 → 不是表格类生成, 不做兜底
            return BlockFillStatus.NEEDS_REVIEW
        data_start = header_idx + 1
        if data_start >= len(table.rows):
            # 无数据行 → 复制表头行作为首数据行
            src = table.rows[header_idx]
            self._clone_table_row(table, src)
        row = table.rows[data_start]
        if row.cells:
            self._set_cell_text(row.cells[0], text)
        return BlockFillStatus.FILLED

    def _fill_response_table(self, doc, block, project, warnings) -> str:
        """生成逐条应答并填充应答表格。

        应答表格结构: 采购文件章节号 | 要求描述 | 响应情况 | 是否偏离 | 偏离描述(模板各异, 按可用列适配)
        """
        from apps.generation.constants import ModelType
        from apps.generation.models import ModelConfig
        from apps.generation.services.llm_service import LLMService

        requirements = self._load_requirements(block, limit=15)
        if not requirements:
            warnings.append(FillWarning(block.block_key, "未找到招标条款, 应答表留空"))
            return BlockFillStatus.NEEDS_REVIEW

        table = self._find_table(doc, block.anchor_text)
        if table is None:
            warnings.append(FillWarning(block.block_key, f"未找到应答表格: {block.anchor_text}"))
            return BlockFillStatus.NEEDS_REVIEW

        req_text = "\n".join(f"[{r['clause_no'] or '?'}] {r['content'][:150]}" for r in requirements)
        system = (
            "你是投标条款应答专家。对招标文件的每一条要求, 逐条生成响应内容。\n"
            "规则:\n"
            "1. 响应情况必须具体、专业, 结合企业能力, 禁止空话; \n"
            "2. 状态只能从 [完全响应, 部分响应, 偏离, 待确认] 中选择, 无法确认的要求标'待确认'; \n"
            "3. 输出 JSON 数组, 每项: {\"clause\": \"章节号\", \"requirement\": \"要求\", \"response\": \"响应内容(80字内)\", \"status\": \"完全响应\", \"deviation\": \"偏离描述, 无则空\"}"
        )
        user = f"招标条款:\n{req_text}\n\n请逐条生成应答。"
        resp = LLMService().chat(
            ModelConfig.objects.get(model_type=ModelType.CHAT, is_default=True, is_active=True),
            system,
            user,
            response_format={"type": "json_object"},
        )
        try:
            data = resp.json if isinstance(resp.json, dict) else json.loads(resp.text)
        except (json.JSONDecodeError, TypeError):
            warnings.append(FillWarning(block.block_key, "应答生成结果非 JSON, 表格留空"))
            return BlockFillStatus.NEEDS_REVIEW
        # 兼容 LLM 直接返回数组的形态
        if isinstance(data, list):
            items = data
        else:
            items = data.get("items") or data.get("responses") or []

        # 写入表格: 表头之后逐行填充; 行不足则复制模板行
        self._write_rows_to_table(table, items, warnings)

        # 风险闸门: 保存应答明细, 标注待确认条目
        review_items = [
            it for it in items if str(it.get("status", "")) == "待确认"
        ]
        block.fill_payload = {
            "items": items,
            "review_count": len(review_items),
        }
        block.save(update_fields=["fill_payload", "updated_at"])
        if review_items:
            warnings.append(FillWarning(
                block.block_key,
                f"应答表含 {len(review_items)} 条'待确认'条目, 生成后需人工复核",
            ))
        return BlockFillStatus.FILLED

    def _write_rows_to_table(self, table, items: list, warnings) -> None:
        header_idx, header_text = self._detect_header_row(table)
        data_rows = table.rows[header_idx + 1:] if header_idx is not None else table.rows[1:]
        n_rows = len(data_rows)

        for i, item in enumerate(items):
            if i >= n_rows:
                new_row = self._clone_table_row(table, data_rows[-1] if data_rows else table.rows[-1])
                data_rows = list(table.rows)[header_idx + 1:]
                n_rows = len(data_rows)
            row = data_rows[i]
            cells = row.cells
            texts = [
                str(item.get("clause", "")),
                str(item.get("requirement", ""))[:80],
                str(item.get("response", "")),
                str(item.get("status", "待确认")),
                str(item.get("deviation", "")),
            ]
            for j, cell in enumerate(cells):
                if j < len(texts) and texts[j]:
                    self._set_cell_text(cell, texts[j])

    def _load_requirements(self, block, limit: int = 15) -> list:
        """从 TenderChunk 加载招标条款(技术/商务/法律要求)。"""
        from apps.tender.constants import ChunkType
        from apps.tender.models import TenderChunk

        template = block.template
        pd = template.parsed_document
        if not pd:
            return []
        chunks = (
            TenderChunk.objects
            .filter(parsed_document=pd, chunk_type__in=[ChunkType.TECH_REQ, ChunkType.COMMERCIAL, ChunkType.LEGAL])
            .order_by("chunk_index")[:limit]
        )
        return [{"clause_no": c.clause_no, "content": c.content} for c in chunks]

    def _build_ai_context(self, block, project, limit: int = 2000) -> str:
        """构建 AI 生成上下文: 项目信息 + 招标条款。"""
        parts = []
        if project:
            parts.append(f"项目名称: {project.name}")
        chunks = self._load_requirements(block, limit=10)
        for c in chunks:
            parts.append(f"[条款 {c['clause_no'] or '?'}] {c['content'][:200]}")
        text = "\n".join(parts)
        return text[:limit]

    # ------------------------------------------------------------------
    # REPEAT_TABLE: 表格行复制 + 案例自动匹配
    # ------------------------------------------------------------------
    # 案例字段 → 表格列关键词映射
    CASE_COLUMN_RULES = [
        ("period", ["起止年月", "起止时间", "实施时间", "项目周期"]),
        ("project_name", ["项目名称"]),
        ("client_name", ["甲方名称", "客户名称", "业主名称"]),
        ("client_contact", ["证明人", "联系人"]),
        ("amount", ["实施金额", "合同金额", "金额"]),
        ("scope", ["范围概述", "项目范围", "范围"]),
        ("remark", ["备注"]),
    ]

    def _repeat_table_rows(self, doc, block, warnings) -> str:
        """复制模板数据行并填充企业案例库数据。"""
        table = self._find_table(doc, block.anchor_text)
        if table is None:
            warnings.append(FillWarning(block.block_key, f"未找到表格: {block.anchor_text}"))
            return BlockFillStatus.NEEDS_REVIEW

        # 数据行 = 不含锚点文本的行(锚点通常在表头)
        anchor_norm = self._normalize(block.anchor_text)
        data_rows = []
        for row in table.rows:
            row_text = self._normalize("".join(c.text for c in row.cells))
            if anchor_norm and anchor_norm in row_text:
                continue  # 表头行
            data_rows.append(row)
        if not data_rows:
            data_rows = [table.rows[-1]]

        repeat_count = int((block.binding_config or {}).get("repeat_count", 3))
        repeat_count = max(1, min(repeat_count, 10))

        # 匹配企业案例
        cases = self._match_cases(block, limit=repeat_count)

        total = max(repeat_count, len(cases))
        template_row = data_rows[0]
        for _ in range(total - 1):
            self._clone_table_row(table, template_row)

        if not cases:
            warnings.append(FillWarning(block.block_key, "企业案例库无匹配案例, 行已复制待人工填写"))
            return BlockFillStatus.NEEDS_REVIEW

        # 按表头识别列, 逐行填充案例
        col_map = self._detect_case_columns(table)
        filled_rows = 0
        for i, case in enumerate(cases):
            row = table.rows[len(table.rows) - total + i]
            for col_idx, attr in col_map.items():
                try:
                    value = getattr(case, attr, "") or ""
                    if callable(value):
                        value = value()
                    if value not in (None, ""):
                        self._set_cell_text(row.cells[col_idx], str(value))
                except Exception:
                    continue
            filled_rows += 1

        block.fill_payload = {
            "cases": [
                {"project_name": c.project_name, "client_name": c.client_name}
                for c in cases
            ],
            "filled": filled_rows,
        }
        block.save(update_fields=["fill_payload", "updated_at"])
        return BlockFillStatus.FILLED

    # 案例匹配关键词停用词(过于泛化, 不参与相关度)
    CASE_KEYWORD_STOPWORDS = {
        "项目", "服务", "采购", "招标", "投标", "公司", "系统",
        "平台", "管理", "建设", "运维", "维护", "评估", "测试", "中心",
    }

    def _match_cases(self, block, limit: int = 5) -> list:
        """从企业案例库匹配案例(v1: 默认企业 + 关键词相关度排序)。"""
        from apps.enterprise.models import CompanyCase, CompanyProfile

        company = CompanyProfile.objects.filter(is_default=True).first()
        qs = CompanyCase.objects.all()
        if company:
            qs = qs.filter(company=company)

        # 关键词: 项目名 jieba 分词(过滤停用词)
        keywords = []
        project = block.template.project
        if project and project.name:
            import jieba

            keywords = [
                k for k in jieba.lcut(project.name)
                if len(k) >= 2 and k not in self.CASE_KEYWORD_STOPWORDS
            ]

        cases = list(qs.order_by("-created_at"))
        if keywords:
            def score(c):
                text = f"{c.project_name} {c.client_name} {c.scope}"
                return sum(1 for k in keywords if k in text)
            cases.sort(key=score, reverse=True)
        return cases[:limit]

    def _detect_case_columns(self, table) -> dict:
        """识别表头各列对应的案例字段。返回 {列index: 字段名}。"""
        header_row = table.rows[0]
        col_map = {}
        for j, cell in enumerate(header_row.cells):
            text = self._normalize(cell.text)
            for attr, keywords in self.CASE_COLUMN_RULES:
                if any(k in text for k in keywords) and attr not in col_map.values():
                    col_map[j] = attr
                    break
        return col_map

    def _clone_table_row(self, table, source_row):
        """deepcopy 行 XML, 追加到表格末尾, 保留全部格式并清空文本。"""
        from copy import deepcopy

        new_tr = deepcopy(source_row._tr)
        # 清空文本
        for t in new_tr.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
            t.text = ""
        source_row._tr.addnext(new_tr)
        return new_tr

    # ------------------------------------------------------------------
    # REPEAT_BLOCK: 整块复制(人员简历等)
    # ------------------------------------------------------------------
    # 块边界标题模式: "一、" "附件N" "第X部分"
    BLOCK_BOUNDARY_RE = re.compile(
        r"^[一二三四五六七八九十]+、|^附件\s*\d+[：:]|^第[一二三四五六七八九十\d]+部分"
    )

    def _repeat_block(self, doc, block, warnings) -> str:
        """定位锚点段落, 复制其后到块边界前的所有元素 N 份(保留格式)。

        锚点段落(块标题)本身保留一份, 复制内容 N-1 份追加在边界前。
        """
        from copy import deepcopy

        anchor = block.anchor_text.strip()
        para = self._locate_paragraph(doc, block)
        if para is None:
            warnings.append(FillWarning(block.block_key, f"未找到锚点段落: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        anchor_el = para._p
        # 收集锚点之后的块元素, 直到遇到边界标题
        block_els = []
        for el in anchor_el.itersiblings():
            text = "".join(el.itertext()).strip()
            if text and self.BLOCK_BOUNDARY_RE.match(text):
                break
            block_els.append(el)
        if not block_els:
            warnings.append(FillWarning(block.block_key, f"锚点后无内容可复制: {anchor}"))
            return BlockFillStatus.NEEDS_REVIEW

        repeat_count = int((block.binding_config or {}).get("repeat_count", 3))
        repeat_count = max(1, min(repeat_count, 10))

        # 复制 N-1 份, 记录每份元素
        copies = []
        insert_after = block_els[-1]
        for _ in range(repeat_count - 1):
            copy_els = []
            for el in block_els:
                new_el = deepcopy(el)
                insert_after.addnext(new_el)
                insert_after = new_el
                copy_els.append(new_el)
            copies.append(copy_els)

        # 人员库填充: 原始块 + 每份复制各匹配一名人员
        members = self._match_members(block, repeat_count)
        filled_members = 0
        all_parts = [block_els] + copies
        for i, part in enumerate(all_parts):
            if i < len(members):
                self._fill_member_fields(part, members[i])
                filled_members += 1

        block.fill_payload = {
            "copied": repeat_count,
            "elements": len(block_els),
            "members_filled": filled_members,
            "note": "块内容已复制, 人员字段已自动填充, 其余空位待人工填写",
        }
        block.save(update_fields=["fill_payload", "updated_at"])
        return BlockFillStatus.FILLED

    # 人员字段 → 锚点关键词
    MEMBER_FIELD_RULES = [
        (re.compile(r"姓名"), "name"),
        (re.compile(r"角色|岗位"), "role"),
        (re.compile(r"工作年限|从业年限|工作经验"), "experience_years"),
        (re.compile(r"职称"), "title"),
        (re.compile(r"证书|资质"), "certificates"),
    ]

    def _match_members(self, block, limit: int = 5) -> list:
        """从人员库匹配人员(v1: 默认企业成员, 按创建顺序取前 N)。"""
        from apps.enterprise.models import CompanyProfile, ProjectMember

        company = CompanyProfile.objects.filter(is_default=True).first()
        qs = ProjectMember.objects.all()
        if company:
            qs = qs.filter(company=company)
        return list(qs.order_by("created_at")[:limit])

    def _fill_member_fields(self, elements: list, member) -> None:
        """在块元素内按关键词定位字段并填充人员数据。"""
        values = {}
        for pattern, attr in self.MEMBER_FIELD_RULES:
            val = getattr(member, attr, None)
            if attr == "experience_years" and val:
                val = f"{val}年"
            if val not in (None, ""):
                values[pattern] = str(val)

        for el in elements:
            for p in el.iter():
                if p.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
                    continue
                text = "".join(p.itertext())
                if not text:
                    continue
                for pattern, val in values.items():
                    if pattern.search(text):
                        new_text = self._replace_first_placeholder(
                            text, val, hint_text=pattern.pattern,
                        )
                        if new_text != text:
                            self._set_lxml_paragraph_text(p, new_text)
                        break

    @staticmethod
    def _set_lxml_paragraph_text(p_el, text: str) -> None:
        """清空段落文本, 写入新文本(保留段落格式)。无 w:t 时创建 run。"""
        from lxml import etree

        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ts = p_el.findall(f".//{{{W_NS}}}t")
        if not ts:
            # 空段落: 创建 run + text
            r = etree.SubElement(p_el, f"{{{W_NS}}}r")
            t = etree.SubElement(r, f"{{{W_NS}}}t")
            t.text = text
            return
        ts[0].text = text
        for t in ts[1:]:
            t.text = ""

    # ------------------------------------------------------------------
    # MATERIAL_SLOT: 材料图片插入
    # ------------------------------------------------------------------
    def _insert_material(self, doc, block, material_package, warnings) -> str:
        usage_key = (block.binding_config or {}).get("usage_key", "")
        material = None
        if material_package and usage_key:
            material = material_package.get_material_by_usage_key(usage_key)
        if material is None or not material.object_key:
            warnings.append(FillWarning(block.block_key, f"缺少材料: {usage_key or '未绑定'}"))
            return BlockFillStatus.NEEDS_REVIEW

        # 图片类型检测
        ext = material.object_key.rsplit(".", 1)[-1].lower() if "." in material.object_key else ""
        if f".{ext}" not in IMAGE_EXTS:
            warnings.append(FillWarning(block.block_key, f"材料非图片, 请人工插入: {material.download_filename()}"))
            return BlockFillStatus.NEEDS_REVIEW

        para = self._locate_paragraph(doc, block)
        if para is None:
            warnings.append(FillWarning(block.block_key, f"未找到材料插入位置: {block.anchor_text}"))
            return BlockFillStatus.NEEDS_REVIEW

        try:
            from docx.shared import Mm

            data = self.storage.get_object(material.object_key)
            run = para.add_run()
            run.add_picture(io.BytesIO(data), width=Mm(140))
            return BlockFillStatus.FILLED
        except Exception as exc:
            logger.exception("material insert failed: key=%s", block.block_key)
            warnings.append(FillWarning(block.block_key, f"材料插入失败: {exc}"))
            return BlockFillStatus.NEEDS_REVIEW

    # ------------------------------------------------------------------
    # 定位工具
    # ------------------------------------------------------------------
    @staticmethod
    def _normalize(text: str) -> str:
        # 全角括号转半角, 消除全半角差异导致的锚点失配
        text = (text or "").replace("（", "(").replace("）", ")")
        return re.sub(r"\s+", "", text)

    def _locate_paragraph(self, doc, block):
        """定位段落: 优先 SDT 控件(Tag=bid.rt:<block_key>), 回退文本锚点(限定附件范围)。"""
        para = self._find_paragraph_by_sdt(doc, block)
        if para is not None:
            return para
        attachment_no = (block.source_config or {}).get("attachment_no")
        return self._find_paragraph(doc, block.anchor_text, attachment_no=attachment_no)

    def _find_paragraph_by_sdt(self, doc, block):
        """按 Content Control Tag 精确定位段落(v2 定位)。"""
        from docx.text.paragraph import Paragraph

        from apps.response_template.services.compile_service import find_sdt_by_tag

        tag = f"bid.rt:{block.block_key}"
        sdt = find_sdt_by_tag(doc.element.body, tag)
        if sdt is None:
            return None
        p_el = sdt.getparent()
        while p_el is not None and p_el.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p":
            p_el = p_el.getparent()
        if p_el is None:
            return None
        # 移除 sdt 标记, 段落内容保留
        sdt.getparent().remove(sdt)
        return Paragraph(p_el, doc)

    def _attachment_range(self, doc, attachment_no):
        """返回指定附件范围内的 body 直接子元素列表(段落+表格)。

        附件边界: 从"附件N"标题到下一个附件标题。
        从"第X部分 ...响应文件格式"章节标题之后开始定位
        (招标文件正文前常有目录/指引里的一套附件标题, 需跳过)。
        """
        if not attachment_no:
            return None
        target = str(attachment_no)
        elements = []
        in_target = False
        body = doc.element.body
        started = False
        for el in body.iterchildren():
            text = "".join(el.itertext()).strip()
            # 定位响应文件格式章节
            if not started:
                if "响应文件格式" in text and re.match(r"^第.+部分", text):
                    started = True
                continue
            m = re.match(r"附件\s*(\d+)", text)
            if m:
                if m.group(1) == target:
                    in_target = True
                elif in_target:
                    break
                continue
            if in_target:
                elements.append(el)
        if not started:
            # 未找到章节标题 → fallback: 从第一个附件标题开始
            in_target = False
            for el in body.iterchildren():
                text = "".join(el.itertext()).strip()
                m = re.match(r"附件\s*(\d+)", text)
                if m:
                    if m.group(1) == target:
                        in_target = True
                    elif in_target:
                        break
                    continue
                if in_target:
                    elements.append(el)
        return elements or None

    def _attachment_range_set(self, doc, attachment_no):
        """附件范围 set(定位加速用)。返回 None 表示不限范围。"""
        if not attachment_no:
            return None
        elements = self._attachment_range(doc, attachment_no)
        if elements is None:
            return None
        return set(elements)

    def _para_in_attachment(self, doc, para, range_set) -> bool:
        """判断段落是否属于附件范围 set。range_set 为 None 时返回 True。"""
        if range_set is None:
            return True
        body = doc.element.body
        el = para._p
        while el.getparent() is not None and el.getparent() is not body:
            el = el.getparent()
        return el in range_set

    def _find_paragraph(self, doc, anchor: str, attachment_no=None):
        """查找含锚点(归一化)的段落。

        匹配优先级(解决歧义):
        1. 归一化完全相等(表格标签如"响应人名称"精确命中, 不误中正文提及)
        2. 含锚点且含空位(表格 cell 优先, 再主文档)
        3. 任意包含(兜底)
        支持 attachment_no 限定搜索范围(同一 label 多附件场景)。
        """
        anchor_norm = self._normalize(anchor)
        if not anchor_norm:
            return None

        range_set = self._attachment_range_set(doc, attachment_no)

        def in_range(para):
            return self._para_in_attachment(doc, para, range_set)

        # 1. 完全相等(表格 + 主文档)
        for para in self._iter_table_paragraphs(doc):
            if not in_range(para):
                continue
            if self._normalize(para.text) == anchor_norm:
                return para
        for para in doc.paragraphs:
            if not in_range(para):
                continue
            if self._normalize(para.text) == anchor_norm:
                return para

        # 2. 包含 + 空位
        for para in self._iter_table_paragraphs(doc):
            if not in_range(para):
                continue
            if anchor_norm in self._normalize(para.text) and self._has_placeholder(para.text):
                return para
        for para in doc.paragraphs:
            if not in_range(para):
                continue
            if anchor_norm in self._normalize(para.text) and self._has_placeholder(para.text):
                return para

        # 3. 任意包含
        for para in self._iter_table_paragraphs(doc):
            if not in_range(para):
                continue
            if anchor_norm in self._normalize(para.text):
                return para
        for para in doc.paragraphs:
            if not in_range(para):
                continue
            if anchor_norm in self._normalize(para.text):
                return para
        return None

    @staticmethod
    def _iter_table_paragraphs(doc):
        """遍历所有表格 cell 内的段落。"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    @staticmethod
    def _has_placeholder(text: str) -> bool:
        return bool(
            UNDERLINE_RE.search(text)
            or PAREN_EMPTY_RE.search(text)
            or PAREN_HINT_RE.search(text)
            or YEAR_MONTH_DAY_RE.search(text)
            or SPACE_RUN_RE.search(text)
        )

    def _find_label_cell(self, doc, anchor: str, attachment_no=None):
        """查找含锚点文本的表格 cell(支持附件范围限定)。"""
        from docx.table import _Cell

        anchor_norm = self._normalize(anchor)
        if not anchor_norm:
            return None
        for table in doc.tables:
            if attachment_no and not self._table_in_attachment(doc, table, attachment_no):
                continue
            for row in table.rows:
                for cell in row.cells:
                    if anchor_norm in self._normalize(cell.text):
                        return cell
        return None

    def _table_in_attachment(self, doc, table, attachment_no) -> bool:
        elements = self._attachment_range(doc, attachment_no)
        if elements is None:
            return True
        return table._tbl in elements

    def _find_table(self, doc, anchor: str):
        """查找含锚点文本(表头/标题)的表格。

        优先单 cell 匹配; 失败后尝试整行拼接匹配(表头跨多 cell 场景,
        如"采购文件章节号 | 要求描述 | 响应人响应情况"整行锚点)。
        """
        anchor_norm = self._normalize(anchor)
        for table in doc.tables:
            if not anchor_norm:
                return table
            # 1. 单 cell 匹配
            for row in table.rows[:4]:
                for cell in row.cells:
                    if anchor_norm in self._normalize(cell.text):
                        return table
            # 2. 整行拼接匹配
            for row in table.rows[:4]:
                row_text = self._normalize("".join(c.text for c in row.cells))
                if anchor_norm in row_text:
                    return table
        return None

    def _detect_header_row(self, table) -> Tuple[Optional[int], str]:
        """识别表头行: 含"要求/响应/偏离/章节"等关键字的行。"""
        keywords = ["要求", "响应", "偏离", "章节", "序号", "项目阶段", "工作项"]
        for i, row in enumerate(table.rows[:3]):
            text = "".join(c.text for c in row.cells)
            if any(k in text for k in keywords):
                return i, text
        return None, ""
