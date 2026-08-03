# backend/apps/knowledge/services/document_parse_service.py
"""知识文档解析服务。"""

import io
import logging

from apps.common.services.storage import StorageService
from apps.knowledge.constants import ParseStatus, ChunkStatus
from apps.knowledge.models import KnowledgeDocument

logger = logging.getLogger(__name__)


class DocumentParseService:
    """知识文档解析服务（复用 tender 解析能力）。"""

    def parse(self, document: KnowledgeDocument) -> None:
        """解析知识文档。

        Args:
            document: 文档实例

        Raises:
            ValueError: 文档尚未准备好解析
        """
        if document.parse_status != ParseStatus.PENDING:
            return

        document.parse_status = ParseStatus.PARSING
        document.save()

        try:
            storage = StorageService()

            # 读取文件内容
            file_content = storage.get_object(document.file_uri)

            # 根据文件类型调用解析
            extension = self._get_extension(document.file_name)
            result = self._do_parse(file_content, document.file_name, extension)

            # 保存解析结果
            parsed_uri = f"knowledge/{document.knowledge_base.id}/{document.id}/parsed.md"
            storage.put_object(parsed_uri, result["markdown"].encode("utf-8"))

            document.parsed_uri = parsed_uri
            document.parser_version = result.get("parser_version", "v1")
            document.metadata["page_count"] = result.get("page_count", 0)
            document.metadata["parse_engine"] = result.get("parse_engine", "builtin")
            document.metadata["table_count"] = result.get("table_count", 0)
            document.metadata["paragraph_count"] = result.get("paragraph_count", 0)

            document.parse_status = ParseStatus.PARSED
            document.chunk_status = ChunkStatus.PENDING
            document.save()

            logger.info(
                "Parsed knowledge document id=%s name=%s engine=%s chars=%d",
                document.id,
                document.file_name,
                result.get("parse_engine"),
                len(result["markdown"]),
            )

        except Exception as e:
            document.parse_status = ParseStatus.FAILED
            document.error_message = str(e)[:2000]
            document.save()
            logger.exception("Failed to parse knowledge document id=%s", document.id)
            raise

    def _get_extension(self, filename: str) -> str:
        """获取文件扩展名（小写）。"""
        if "." in filename:
            return filename.rsplit(".", 1)[-1].lower()
        return ""

    def _do_parse(self, content: bytes, filename: str, extension: str) -> dict:
        """执行解析。

        Args:
            content: 文件二进制内容
            filename: 文件名
            extension: 文件扩展名

        Returns:
            解析结果字典
        """
        if extension == "docx":
            return self._parse_word(content, filename)
        elif extension == "doc":
            from apps.common.services.doc_converter import DocConverter

            docx_content = DocConverter().convert_doc_to_docx(content, filename)
            return self._parse_word(docx_content, filename)
        elif extension == "pdf":
            return self._parse_pdf(content, filename)
        elif extension in ["txt", "md", "markdown"]:
            return self._parse_text(content)
        else:
            raise ValueError(f"不支持的文件格式: {extension}")

    def _parse_word(self, content: bytes, file_name: str) -> dict:
        """解析 Word 文档（使用 python-docx）。"""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安装，请添加到 requirements.txt")

        document = Document(io.BytesIO(content))

        markdown_parts = []
        table_count = 0
        paragraph_count = 0

        # 遍历文档元素
        for element in document.element.body:
            # 处理段落
            if element.tag.endswith("}p"):
                paragraph = self._find_paragraph_by_element(document, element)
                if paragraph:
                    text = self._process_paragraph(paragraph)
                    if text:
                        markdown_parts.append(text)
                        paragraph_count += 1

            # 处理表格
            elif element.tag.endswith("}tbl"):
                table = self._find_table_by_element(document, element)
                if table:
                    table_md = self._process_table(table)
                    if table_md:
                        markdown_parts.append(table_md)
                        table_count += 1

        # 处理图片（提取图片信息）
        image_count = 0
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
                # P0: 只记录图片数量，不提取内容
                # P1: 可接入 OCR 提取图片文字

        if image_count > 0:
            markdown_parts.append(f"\n\n**文档包含 {image_count} 张图片**")

        markdown = "\n\n".join(markdown_parts)

        return {
            "markdown": markdown,
            "page_count": 0,
            "table_count": table_count,
            "paragraph_count": paragraph_count,
            "parse_engine": "docx-python-docx",
            "parser_version": "v2",
        }

    def _parse_pdf(self, content: bytes, file_name: str) -> dict:
        """解析 PDF 文件。

        使用 pdfplumber 提取文本与表格。pdfplumber 未安装时抛异常，
        让 pipeline 标记文档 FAILED，避免占位文本污染 RAG 检索。
        """
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError("PDF 解析依赖未安装：pdfplumber，请更新 requirements.txt 并重建镜像")

        markdown_parts = []
        page_count = 0

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, start=1):
                # 提取文本
                text = page.extract_text()
                if text:
                    # 添加页码标记
                    markdown_parts.append(f"## 第 {page_num} 页\n\n{text}")

                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        table_md = self._format_table(table)
                        markdown_parts.append(table_md)

        markdown = "\n\n".join(markdown_parts)

        return {
            "markdown": markdown,
            "page_count": page_count,
            "parse_engine": "pdfplumber",
            "parser_version": "v2",
        }

    def _parse_text(self, content: bytes) -> dict:
        """解析纯文本/Markdown 文件。"""
        # 处理编码异常
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("gbk", errors="ignore")

        return {
            "markdown": text,
            "page_count": 1,
            "parse_engine": "text",
            "parser_version": "v1",
        }

    def _find_paragraph_by_element(self, document, element):
        """根据 XML 元素找到对应的 Paragraph 对象。"""
        from docx.text.paragraph import Paragraph
        for para in document.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_by_element(self, document, element):
        """根据 XML 元素找到对应的 Table 对象。"""
        from docx.table import Table
        for table in document.tables:
            if table._element is element:
                return table
        return None

    def _process_paragraph(self, paragraph) -> str:
        """处理段落，返回 Markdown 文本。"""
        text = paragraph.text.strip()
        if not text:
            return ""

        # 根据样式判断标题级别
        style_name = paragraph.style.name if paragraph.style else ""

        if "Heading 1" in style_name or "标题 1" in style_name:
            return f"# {text}"
        elif "Heading 2" in style_name or "标题 2" in style_name:
            return f"## {text}"
        elif "Heading 3" in style_name or "标题 3" in style_name:
            return f"### {text}"
        elif "Heading 4" in style_name or "标题 4" in style_name:
            return f"#### {text}"
        elif "Heading 5" in style_name or "标题 5" in style_name:
            return f"##### {text}"
        elif "Heading 6" in style_name or "标题 6" in style_name:
            return f"###### {text}"
        elif "List" in style_name or "列表" in style_name:
            return f"- {text}"

        # 检测可能的标题格式
        import re
        if re.match(r"^[一二三四五六七八九十]+、", text):
            return f"# {text}"
        if re.match(r"^\d+\.\s+\S", text):
            if len(text) < 50:
                return f"## {text}"

        return text

    def _process_table(self, table) -> str:
        """处理表格，转换为 Markdown 格式。"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    para.text.strip() for para in cell.paragraphs if para.text.strip()
                )
                cell_text = cell_text.replace("|", "\\|")
                cells.append(cell_text)
            rows.append(cells)

        if not rows:
            return ""

        lines = []

        # 表头
        if rows:
            header = rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 数据行
        for row in rows[1:]:
            while len(row) < len(rows[0]):
                row.append("")
            lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

        return "\n".join(lines)

    def _format_table(self, table: list) -> str:
        """格式化 PDF 提取的表格为 Markdown。"""
        if not table or not table[0]:
            return ""

        lines = []

        # 表头
        header = table[0]
        lines.append("| " + " | ".join(str(cell or "") for cell in header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 数据行
        for row in table[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(str(cell or "") for cell in row[: len(header)]) + " |")

        return "\n".join(lines)