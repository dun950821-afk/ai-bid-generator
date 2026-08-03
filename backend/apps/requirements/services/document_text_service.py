# backend/apps/requirements/services/document_text_service.py
"""文档全文提取服务。

从 TenderFile 中提取纯文本全文，存储到 MinIO，用于条款抽取。
独立于 TenderChunk 分块逻辑。
"""

import hashlib
import io
import logging

from apps.common.services.storage import StorageService
from apps.tender.models import TenderFile

logger = logging.getLogger(__name__)


class DocumentTextService:
    """文档全文提取服务。

    从 DOCX 文件中提取纯文本，存储到 MinIO。
    文本不入数据库，仅存 MinIO，通过 object_key 引用。
    """

    def __init__(self, storage: StorageService | None = None):
        self.storage = storage or StorageService()

    def extract_and_store(self, tender_file: TenderFile) -> tuple[str, str]:
        """提取文档全文并存储到 MinIO。

        Args:
            tender_file: TenderFile 实例

        Returns:
            (object_key, text_hash) 元组

        Raises:
            ValueError: 文件格式不支持
            StorageError: 存储失败
        """
        # 检查缓存：如果已有且文件未变，直接返回
        if tender_file.document_text_object_key and tender_file.document_text_hash:
            if self.storage.object_exists(tender_file.document_text_object_key):
                logger.info(
                    "Document text already extracted: tender_file=%s key=%s",
                    tender_file.id,
                    tender_file.document_text_object_key,
                )
                return tender_file.document_text_object_key, tender_file.document_text_hash

        # 读取原始文件
        content = self.storage.get_object(tender_file.object_key)

        # 提取文本
        text = self._extract_text(content, tender_file.original_name)

        # 计算哈希
        text_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()

        # 生成存储路径
        object_key = self._build_object_key(tender_file)

        # 存储到 MinIO
        self.storage.put_object(
            object_key,
            text.encode("utf-8"),
            content_type="text/plain; charset=utf-8",
        )

        # 更新 TenderFile 记录
        tender_file.document_text_object_key = object_key
        tender_file.document_text_hash = text_hash
        tender_file.save(update_fields=["document_text_object_key", "document_text_hash", "updated_at"])

        logger.info(
            "Document text extracted and stored: tender_file=%s key=%s chars=%d",
            tender_file.id,
            object_key,
            len(text),
        )

        return object_key, text_hash

    def get_document_text(self, tender_file: TenderFile) -> str:
        """获取文档全文。

        如果尚未提取，则自动提取并存储。

        Args:
            tender_file: TenderFile 实例

        Returns:
            文档全文

        Raises:
            ValueError: 文件格式不支持
            StorageError: 存储失败
        """
        # 检查是否有缓存
        if tender_file.document_text_object_key:
            try:
                content = self.storage.get_object(tender_file.document_text_object_key)
                return content.decode("utf-8")
            except Exception:
                logger.warning(
                    "Document text cache miss, re-extracting: tender_file=%s",
                    tender_file.id,
                )

        # 提取并存储
        object_key, _ = self.extract_and_store(tender_file)
        content = self.storage.get_object(object_key)
        return content.decode("utf-8")

    def _extract_text(self, content: bytes, filename: str) -> str:
        """从文件内容中提取纯文本。

        Args:
            content: 文件二进制内容
            filename: 文件名（用于判断格式）

        Returns:
            提取的纯文本
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext == "docx":
            return self._extract_from_docx(content)
        elif ext == "pdf":
            return self._extract_from_pdf(content)
        elif ext in ("txt", "md"):
            return content.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _extract_from_docx(self, content: bytes) -> str:
        """从 DOCX 文件中提取纯文本。

        Args:
            content: DOCX 文件二进制内容

        Returns:
            提取的纯文本
        """
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安装，请添加到 requirements.txt")

        document = Document(io.BytesIO(content))

        text_parts = []

        # 提取所有段落文本
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # 提取表格文本
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        # 用换行符连接所有文本
        full_text = "\n\n".join(text_parts)

        return full_text

    def _extract_from_pdf(self, content: bytes) -> str:
        """从 PDF 文件中提取纯文本。

        Args:
            content: PDF 文件二进制内容

        Returns:
            提取的纯文本
        """
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError("pdfplumber 未安装，请添加到 requirements.txt")

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                # 提取表格文本（与 DOCX 表格同风格：单元格以 | 分隔）
                for table in page.extract_tables() or []:
                    for row in table:
                        row_text = []
                        for cell in row:
                            cell_text = (cell or "").strip().replace("\n", " ")
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                # 页面纯文本
                text = page.extract_text() or ""
                if text:
                    text_parts.append(text)

        return "\n\n".join(text_parts)

    def _build_object_key(self, tender_file: TenderFile) -> str:
        """构建文档全文的存储路径。

        Args:
            tender_file: TenderFile 实例

        Returns:
            MinIO 对象键
        """
        return f"parsed/{tender_file.id}/document_text.txt"
