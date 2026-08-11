"""模板渲染服务端到端测试（Phase 3）。"""
import zipfile
from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from django.core.cache import cache
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from rest_framework.test import APIClient

from apps.common.services.storage import ObjectNotFound, StorageService
from apps.enterprise.models import BidMaterialPackage, CompanyMaterial, CompanyProfile
from apps.outline.models import (
    BidWordTemplate,
    BidWordTemplateVersion,
    Outline,
    Section,
)
from apps.outline.services.template.template_render_service import (
    render_bid_document,
)
from apps.projects.models import Lot, Project

User = get_user_model()

# 1x1 透明 PNG
PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _sdt_xml(tag: str, alias: str, text: str, sdt_id: int) -> str:
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def _make_template_docx(controls) -> bytes:
    doc = Document()
    for i, (tag, alias, text) in enumerate(controls):
        p = doc.add_paragraph()
        p._p.append(parse_xml(_sdt_xml(tag, alias, text, 2000 + i)))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _patch_storage(monkeypatch):
    store = {}

    def fake_put(self, key, data, content_type="application/octet-stream"):
        store[key] = data

    def fake_get(self, key):
        if key not in store:
            raise ObjectNotFound(key)
        return store[key]

    monkeypatch.setattr(StorageService, "put_object", fake_put)
    monkeypatch.setattr(StorageService, "get_object", fake_get)
    return store


@pytest.mark.django_db
class TestTemplateRenderService:
    def setup_method(self):
        cache.clear()
        self.user = User.objects.create_user(username="render-user", password="pass")
        self.project = Project.objects.create(name="测试项目", created_by=self.user)
        self.lot = Lot.objects.create(name="标段一", code="01", project=self.project)
        self.outline = Outline.objects.create(
            project=self.project, lot=self.lot, name="测试标书",
            source="preset", created_by=self.user,
        )
        Section.objects.create(
            outline=self.outline, title="项目概述", level=1, sort_order=0,
            content="本项目为<银行>&核心系统\"改造\"项目\n\n- 需求一\n- 需求二",
        )
        Section.objects.create(
            outline=self.outline, title="技术方案", level=1, sort_order=1,
            content="| 项 | 值 |\n|---|---|\n| 工期 | 30天 |",
        )
        self.company = CompanyProfile.objects.create(
            name="测试<科技>&有限公司", legal_representative="李四",
            created_by=self.user,
        )
        self.package = BidMaterialPackage.objects.create(
            outline=self.outline, company=self.company, name="材料包",
        )

    def _make_template(self, monkeypatch, controls, style_mapping=None):
        store = _patch_storage(monkeypatch)
        content = _make_template_docx(controls)
        template = BidWordTemplate.objects.create(
            name="渲染模板", code=f"render-tpl-{id(controls)}", scope_type="system",
            status="active",
        )
        version = BidWordTemplateVersion.objects.create(
            template=template, version_no=1,
            object_key=f"bid-templates/system/{template.id}/versions/v1.docx",
            file_name="v1.docx", file_size=len(content),
            file_hash=f"hash-{template.id}",
            style_mapping=style_mapping or {},
        )
        template.published_version = version
        template.save(update_fields=["published_version"])
        store[version.object_key] = content
        return template, version, store

    def test_render_end_to_end(self, monkeypatch):
        """变量替换 + 正文插入 + 快照落库 + 使用计数。"""
        template, version, store = self._make_template(monkeypatch, [
            ("bid.var:project.name", "项目名称", "项目名称"),
            ("bid.var:company.name", "企业名称", "企业名称"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])

        document, warnings = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )

        # BidDocument 落库 + 快照
        assert document.id is not None
        assert document.template_id == template.id
        assert document.template_version_id == version.id
        assert document.template_file_hash == version.file_hash
        assert document.render_context_snapshot["company"]["name"] == "测试<科技>&有限公司"
        assert document.object_key in store

        # 渲染产物内容
        text = _docx_text(store[document.object_key])
        assert "测试项目" in text  # project.name 变量替换
        assert "测试<科技>&有限公司" in text  # autoescape 后仍还原为原文
        assert "项目概述" in text  # 章节标题
        assert "需求一" in text  # 列表
        assert "30天" in text  # 表格
        assert "{{" not in text  # 无残留指令

        template.refresh_from_db()
        assert template.usage_count == 1

    def test_render_special_chars_autoescape(self, monkeypatch):
        """< > & 等特殊字符不破坏文档（autoescape 生效）。"""
        template, version, store = self._make_template(monkeypatch, [
            ("bid.var:company.name", "企业名称", "企业名称"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        # 能打开即说明 XML 未被特殊字符破坏
        text = _docx_text(store[document.object_key])
        assert "测试<科技>&有限公司" in text

    def test_render_material_image(self, monkeypatch):
        """bid.material 控件解析为材料包中的图片。"""
        material = CompanyMaterial.objects.create(
            company=self.company, material_type="qualification",
            title="营业执照", object_key="materials/license.png",
        )
        from apps.enterprise.models import BidMaterialPackageItem

        BidMaterialPackageItem.objects.create(
            package=self.package, material=material, usage_key="business_license",
        )
        template, version, store = self._make_template(monkeypatch, [
            ("bid.material:business_license", "营业执照", "营业执照"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        store["materials/license.png"] = PNG_1X1

        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        # 图片嵌入后 docx 会包含 media part
        with zipfile.ZipFile(BytesIO(store[document.object_key])) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
        assert media, "材料图片应嵌入渲染产物"

    def test_render_missing_material_shows_placeholder(self, monkeypatch):
        template, version, store = self._make_template(monkeypatch, [
            ("bid.material:not_exist", "不存在", "不存在"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        text = _docx_text(store[document.object_key])
        assert "【缺少材料：not_exist】" in text

    def test_render_with_style_mapping_fallback(self, monkeypatch):
        """样式映射指向不存在的样式时静默降级，不中断生成。"""
        template, version, store = self._make_template(
            monkeypatch,
            [
                ("bid.slot:body", "标书正文", "标书正文"),
            ],
            style_mapping={"heading1": "不存在的样式"},
        )
        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        text = _docx_text(store[document.object_key])
        assert "项目概述" in text

    def test_render_empty_content_warning(self, monkeypatch):
        empty_lot = Lot.objects.create(name="标段二", project=self.project)
        empty_outline = Outline.objects.create(
            project=self.project, lot=empty_lot, name="空标书",
            source="preset", created_by=self.user,
        )
        Section.objects.create(
            outline=empty_outline, title="空章节", level=1, sort_order=0, content="",
        )
        template, version, _ = self._make_template(monkeypatch, [
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        _, warnings = render_bid_document(
            template=template, version=version,
            outline=empty_outline, user=self.user,
        )
        assert any(w["type"] == "no_content" for w in warnings)

    def test_compile_cache_used(self, monkeypatch):
        """同一 file_hash 第二次渲染命中编译缓存（方案 §55）。"""
        template, version, _ = self._make_template(monkeypatch, [
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        assert cache.get(f"tpl-compiled:{version.file_hash}") is not None


@pytest.mark.django_db
class TestBuildDocxWithTemplateAPI:
    def setup_method(self):
        cache.clear()
        self.user = User.objects.create_user(username="api-user", password="pass")
        self.project = Project.objects.create(name="API项目", created_by=self.user)
        self.lot = Lot.objects.create(name="标段", project=self.project)
        role_needed = None
        from apps.projects.models import ProjectMember, ProjectRole

        role = ProjectRole.objects.create(
            project=self.project, name="负责人", code="owner",
            permissions=["project.view", "outline.view", "outline.edit",
                         "export.create"],
            is_builtin=True,
        )
        ProjectMember.objects.create(
            project=self.project, user=self.user, project_role=role,
        )
        self.outline = Outline.objects.create(
            project=self.project, lot=self.lot, name="API标书",
            source="preset", created_by=self.user,
        )
        Section.objects.create(
            outline=self.outline, title="章节", level=1, sort_order=0,
            content="内容",
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)

    def _make_template(self, monkeypatch, published=True):
        store = _patch_storage(monkeypatch)
        content = _make_template_docx([
            ("bid.var:project.name", "项目名称", "项目名称"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        template = BidWordTemplate.objects.create(
            name="API模板", code="api-tpl", scope_type="system", status="active",
        )
        if published:
            version = BidWordTemplateVersion.objects.create(
                template=template, version_no=1,
                object_key=f"bid-templates/system/{template.id}/versions/v1.docx",
                file_name="v1.docx", file_size=len(content),
                file_hash=f"api-hash-{template.id}",
            )
            template.published_version = version
            template.save(update_fields=["published_version"])
            store[version.object_key] = content
        return template

    def test_build_docx_with_template(self, monkeypatch):
        template = self._make_template(monkeypatch)
        resp = self.client.post(
            f"/api/outlines/{self.outline.id}/build_docx/",
            {"template_id": template.id},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["template"]["id"] == template.id
        assert data["template"]["version_no"] == 1

        from apps.outline.models import BidDocument

        document = BidDocument.objects.get(id=data["document_id"])
        assert document.template_id == template.id
        assert document.object_key

    def test_build_docx_template_not_published(self, monkeypatch):
        template = self._make_template(monkeypatch, published=False)
        resp = self.client.post(
            f"/api/outlines/{self.outline.id}/build_docx/",
            {"template_id": template.id},
        )
        assert resp.status_code == 400
        assert resp.json()["code"] == "TEMPLATE_NOT_PUBLISHED"

    def test_build_docx_template_not_found(self):
        resp = self.client.post(
            f"/api/outlines/{self.outline.id}/build_docx/",
            {"template_id": 999999},
        )
        assert resp.status_code == 404
        assert resp.json()["code"] == "TEMPLATE_NOT_FOUND"

    def test_build_docx_without_template_still_works(self, monkeypatch):
        """旧裸生成链路保持可用（过渡期）。"""
        _patch_storage(monkeypatch)
        resp = self.client.post(f"/api/outlines/{self.outline.id}/build_docx/")
        assert resp.status_code == 200
        assert "template" not in resp.json()
