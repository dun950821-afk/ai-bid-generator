"""Markdown → AST 解析器与 AST 渲染器测试（Phase 4）。"""
import zipfile
from io import BytesIO

import pytest
from docx import Document

from apps.common.services.storage import ObjectNotFound, StorageService
from apps.outline.services.document.ast import (
    HeadingNode,
    ImageNode,
    ListNode,
    MaterialNode,
    PageBreakNode,
    ParagraphNode,
    QuoteNode,
    TableNode,
)
from apps.outline.services.document.image_resolver import extract_object_key
from apps.outline.services.document.markdown_parser import MarkdownParser
from apps.outline.services.document.word_body_renderer import WordBodyRenderer

PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


class TestMarkdownParser:
    def setup_method(self):
        self.parser = MarkdownParser()

    def test_headings(self):
        nodes = self.parser.parse("# 一级\n## 二级\n### 三级")
        assert [n.level for n in nodes] == [1, 2, 3]
        assert all(isinstance(n, HeadingNode) for n in nodes)

    def test_bold_and_italic_segments(self):
        nodes = self.parser.parse("普通**粗体**和*斜体*文本")
        assert len(nodes) == 1
        segments = nodes[0].segments
        bold = [s for s in segments if s.bold]
        italic = [s for s in segments if s.italic]
        assert bold and bold[0].text == "粗体"
        assert italic and italic[0].text == "斜体"

    def test_unordered_list_grouped(self):
        nodes = self.parser.parse("- 甲\n- 乙\n* 丙")
        assert len(nodes) == 1
        node = nodes[0]
        assert isinstance(node, ListNode) and not node.ordered
        assert len(node.items) == 3

    def test_ordered_list(self):
        nodes = self.parser.parse("1. 第一\n2. 第二")
        node = nodes[0]
        assert isinstance(node, ListNode) and node.ordered
        assert node.items[0][0].text == "第一"

    def test_table_with_header(self):
        nodes = self.parser.parse("| 列A | 列B |\n|---|---|\n| 1 | 2 |")
        node = nodes[0]
        assert isinstance(node, TableNode)
        assert node.has_header
        assert node.rows == [["列A", "列B"], ["1", "2"]]

    def test_image_standalone(self):
        nodes = self.parser.parse("![架构图](/minio/bucket/editor/images/a.png)")
        assert len(nodes) == 1
        node = nodes[0]
        assert isinstance(node, ImageNode)
        assert node.alt == "架构图"
        assert node.url == "/minio/bucket/editor/images/a.png"

    def test_html_img_with_width(self):
        """原生 HTML img（带 width 属性时 Turndown 保留 HTML 原样）。"""
        nodes = self.parser.parse(
            '<img src="/minio/bid-files/editor/images/2026/08/11/x.png" width="244" alt="">'
        )
        assert len(nodes) == 1
        node = nodes[0]
        assert isinstance(node, ImageNode)
        assert node.url == "/minio/bid-files/editor/images/2026/08/11/x.png"
        assert node.width_px == 244

    def test_html_img_mixed_with_text(self):
        nodes = self.parser.parse(
            '前文 <img src="/minio/bid-files/editor/images/y.png"> 后文'
        )
        kinds = [type(n).__name__ for n in nodes]
        assert "ParagraphNode" in kinds
        assert "ImageNode" in kinds

    def test_material_placeholder(self):
        nodes = self.parser.parse("{{ material:business_license }}")
        node = nodes[0]
        assert isinstance(node, MaterialNode)
        assert node.usage_key == "business_license"
        assert node.is_attachment

    def test_page_break(self):
        nodes = self.parser.parse("上文\n\n---\n\n下文")
        assert any(isinstance(n, PageBreakNode) for n in nodes)

    def test_quote(self):
        nodes = self.parser.parse("> 引用内容")
        assert isinstance(nodes[0], QuoteNode)

    def test_empty_content(self):
        assert self.parser.parse("") == []
        assert self.parser.parse(None) == []


class TestImageResolver:
    def test_proxy_url(self, settings):
        settings.MINIO_BUCKET = "bid"
        assert (
            extract_object_key("/minio/bid/editor/images/2026/08/11/x.png")
            == "editor/images/2026/08/11/x.png"
        )

    def test_direct_url(self, settings):
        settings.MINIO_BUCKET = "bid"
        assert (
            extract_object_key("http://localhost:9000/bid/editor/images/x.png?sign=abc")
            == "editor/images/x.png"
        )

    def test_external_url_rejected(self, settings):
        settings.MINIO_BUCKET = "bid"
        assert extract_object_key("https://evil.com/other/x.png") is None
        assert extract_object_key("") is None


def _render_to_docx(sections, monkeypatch=None, storage_map=None):
    """把假章节渲染进空白 Document，返回 (docx_bytes, store)。"""
    if storage_map is not None and monkeypatch is not None:
        def fake_get(self, key):
            if key not in storage_map:
                raise ObjectNotFound(key)
            return storage_map[key]

        monkeypatch.setattr(StorageService, "get_object", fake_get)

    class FakeSection:
        def __init__(self, id, title, content, sort_order=0, parent_id=None, level=1):
            self.id = id
            self.title = title
            self.content = content
            self.sort_order = sort_order
            self.parent_id = parent_id
            self.level = level

    doc = Document()
    WordBodyRenderer().render(doc, sections)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestWordBodyRendererAst:
    def test_bold_run_in_output(self):
        content = _render_to_docx([
            type("S", (), {"id": 1, "title": "章节", "sort_order": 0,
                           "parent_id": None, "level": 1,
                           "content": "这是**重点**内容"})()
        ])
        doc = Document(BytesIO(content))
        bold_texts = [
            run.text
            for p in doc.paragraphs
            for run in p.runs
            if run.bold
        ]
        assert "重点" in bold_texts

    def test_image_embedded(self, monkeypatch, settings):
        png_key = "editor/images/2026/08/11/x.png"
        image_url = f"/minio/{settings.MINIO_BUCKET}/{png_key}"
        content = _render_to_docx(
            [
                type("S", (), {"id": 1, "title": "章节", "sort_order": 0,
                               "parent_id": None, "level": 1,
                               "content": f"![系统架构图]({image_url})"})()
            ],
            monkeypatch=monkeypatch,
            storage_map={png_key: PNG_1X1},
        )
        with zipfile.ZipFile(BytesIO(content)) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
        assert media, "图片应嵌入 docx"
        doc = Document(BytesIO(content))
        assert any("系统架构图" in p.text for p in doc.paragraphs)

    def test_external_image_placeholder(self):
        content = _render_to_docx([
            type("S", (), {"id": 1, "title": "章节", "sort_order": 0,
                           "parent_id": None, "level": 1,
                           "content": "![x](https://evil.com/a.png)"})()
        ])
        doc = Document(BytesIO(content))
        assert any("图片无法解析" in p.text for p in doc.paragraphs)

    def test_page_break_written(self):
        content = _render_to_docx([
            type("S", (), {"id": 1, "title": "章节", "sort_order": 0,
                           "parent_id": None, "level": 1,
                           "content": "上文\n\n---\n\n下文"})()
        ])
        xml = zipfile.ZipFile(BytesIO(content)).read("word/document.xml").decode()
        assert 'w:type="page"' in xml

    def test_table_header_bold(self):
        content = _render_to_docx([
            type("S", (), {"id": 1, "title": "章节", "sort_order": 0,
                           "parent_id": None, "level": 1,
                           "content": "| 列A | 列B |\n|---|---|\n| 1 | 2 |"})()
        ])
        doc = Document(BytesIO(content))
        header_cell = doc.tables[0].rows[0].cells[0]
        assert all(run.bold for run in header_cell.paragraphs[0].runs if run.text)

    def test_indent_only_body_inherits_table_cells_cleared(self, monkeypatch, settings):
        """Normal 带首行缩进时：正文段落继承缩进，表格单元格/图片显式清零。"""
        import zipfile as zf_mod

        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        png_key = "editor/images/x.png"
        image_url = f"/minio/{settings.MINIO_BUCKET}/{png_key}"

        def fake_get(self, key):
            if key != png_key:
                raise ObjectNotFound(key)
            return PNG_1X1

        monkeypatch.setattr(StorageService, "get_object", fake_get)

        doc = Document()
        # 模拟模板：Normal 首行缩进 2 字符
        ppr = doc.styles["Normal"].element.get_or_add_pPr()
        ppr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLineChars="200"/>'))

        section = type("S", (), {
            "id": 1, "title": "章节", "sort_order": 0, "parent_id": None, "level": 1,
            "content": (
                "正文段落\n\n| 列A |\n|---|\n| 值 |\n\n"
                f"![图]({image_url})"
            ),
        })()
        WordBodyRenderer().render(doc, [section])

        # 正文段落：无直接 ind → 继承 Normal 的 200
        body_p = next(p for p in doc.paragraphs if p.text == "正文段落")
        assert body_p._p.find(qn("w:pPr")) is None or body_p._p.find(
            qn("w:pPr")
        ).find(qn("w:ind")) is None

        # 表格单元格：显式 firstLineChars=0
        cell_p = doc.tables[0].rows[1].cells[0].paragraphs[0]
        ind = cell_p._p.find(qn("w:pPr")).find(qn("w:ind"))
        assert ind is not None and ind.get(qn("w:firstLineChars")) == "0"

        # 图片段落：显式 firstLineChars=0
        buffer = BytesIO()
        doc.save(buffer)
        xml = zf_mod.ZipFile(BytesIO(buffer.getvalue())).read("word/document.xml").decode()
        assert xml.count('w:firstLineChars="0"') >= 2
