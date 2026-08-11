"""补齐功能测试：多册分槽渲染 / 页眉 Logo 替换 / 招标方变量。"""
import zipfile
from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm

from apps.common.services.storage import ObjectNotFound, StorageService
from apps.outline.models import (
    BidWordTemplate,
    BidWordTemplateVersion,
    Outline,
    Section,
)
from apps.outline.services.document.word_post_processor import (
    replace_tagged_images,
)
from apps.outline.services.template.template_render_service import (
    render_bid_document,
)
from apps.outline.services.template.template_validator import TemplateValidator
from apps.projects.models import Lot, Project

User = get_user_model()

PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)
# 替换目标（无需是合法图片——只验证 zip 内字节被覆盖）
LOGO_BYTES = b"LOGO-REPLACED-BYTES"


def _sdt_xml(tag, alias, text, sdt_id):
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def _make_docx(controls) -> bytes:
    doc = Document()
    for i, (tag, alias, text) in enumerate(controls):
        p = doc.add_paragraph()
        p._p.append(parse_xml(_sdt_xml(tag, alias, text, 6000 + i)))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def _patch_storage(monkeypatch):
    store = {}

    def fake_put(self, key, data, content_type="application/octet-stream"):
        store[key] = data

    def fake_get(self, key):
        if key not in store:
            raise ObjectNotFound(key)
        return store[key]

    monkeypatch.setattr(StorageService, "put_object", fake_put)
    monkeypatch.setattr(StorageService, "get_object", fake_get)
    return store


# ---------- 多册分槽 ----------

@pytest.mark.django_db
class TestMultiSlotRender:
    def setup_method(self):
        self.user = User.objects.create_user(username="ms-user", password="pass")
        self.project = Project.objects.create(name="分册项目", created_by=self.user)
        self.lot = Lot.objects.create(name="标段", project=self.project)
        self.outline = Outline.objects.create(
            project=self.project, lot=self.lot, name="分册标书",
            source="preset", created_by=self.user,
        )
        Section.objects.create(
            outline=self.outline, title="技术方案", level=1, sort_order=0,
            content="技术方案正文",
            content_matrix={"section_role": "technical_solution"},
        )
        Section.objects.create(
            outline=self.outline, title="商务响应", level=1, sort_order=1,
            content="商务响应正文",
            content_matrix={"section_role": "business_response"},
        )

    def _make_template(self, monkeypatch, controls):
        store = _patch_storage(monkeypatch)
        content = _make_docx(controls)
        template = BidWordTemplate.objects.create(
            name="分册模板", code=f"ms-tpl-{id(controls)}", scope_type="system",
        )
        version = BidWordTemplateVersion.objects.create(
            template=template, version_no=1,
            object_key=f"bid-templates/system/{template.id}/versions/v1.docx",
            file_name="v1.docx", file_size=len(content),
            file_hash=f"ms-hash-{template.id}",
        )
        store[version.object_key] = content
        return template, version, store

    def test_role_slot_only_renders_matching_sections(self, monkeypatch):
        """技术册插槽只渲染 technical_solution 角色章节。"""
        template, version, store = self._make_template(monkeypatch, [
            ("bid.slot:role.technical_solution", "技术册", "技术册"),
        ])
        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        text = _docx_text(store[document.object_key])
        assert "技术方案正文" in text
        assert "商务响应正文" not in text

    def test_multiple_role_slots(self, monkeypatch):
        """两个分册插槽各自渲染对应章节。"""
        template, version, store = self._make_template(monkeypatch, [
            ("bid.slot:role.technical_solution", "技术册", "技术册"),
            ("bid.slot:role.business_response", "商务册", "商务册"),
        ])
        document, _ = render_bid_document(
            template=template, version=version,
            outline=self.outline, user=self.user,
        )
        text = _docx_text(store[document.object_key])
        assert "技术方案正文" in text
        assert "商务响应正文" in text

    def test_validator_accepts_role_slot_without_body(self):
        content = _make_docx([
            ("bid.slot:role.technical_solution", "技术册", "技术册"),
        ])
        result = TemplateValidator().validate(content)
        assert result["valid"] is True, result["errors"]

    def test_validator_rejects_unknown_role_slot(self):
        content = _make_docx([
            ("bid.slot:role.not_exist_role", "不存在", "不存在"),
        ])
        result = TemplateValidator().validate(content)
        assert result["valid"] is False
        assert any(e["code"] == "VARIABLE_UNKNOWN" for e in result["errors"])

    def test_validator_rejects_duplicate_role_slot(self):
        content = _make_docx([
            ("bid.slot:role.technical_solution", "技术册", "技术册"),
            ("bid.slot:role.technical_solution", "技术册2", "技术册2"),
        ])
        result = TemplateValidator().validate(content)
        assert result["valid"] is False
        assert any(e["code"] == "BODY_SLOT_DUPLICATED" for e in result["errors"])


# ---------- 页眉 Logo 替换 ----------

def _make_docx_with_header_logo(descr="bid.image:company.logo") -> tuple[bytes, str]:
    """构造页眉带标记占位图的 docx，返回 (字节, 媒体 part 名)。"""
    doc = Document()
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(BytesIO(PNG_1X1), width=Mm(20))

    inline = paragraph._p.find(".//" + qn("wp:inline"))
    doc_pr = inline.find(qn("wp:docPr"))
    doc_pr.set("descr", descr)

    doc.add_paragraph("正文")
    buffer = BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    with zipfile.ZipFile(BytesIO(content)) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    return content, media[0]


class TestHeaderImageReplacement:
    def test_tagged_image_replaced(self):
        content, media_part = _make_docx_with_header_logo()
        result = replace_tagged_images(
            content, {"bid.image:company.logo": LOGO_BYTES}
        )
        with zipfile.ZipFile(BytesIO(result)) as zf:
            assert zf.read(media_part) == LOGO_BYTES

    def test_untagged_image_untouched(self):
        content, media_part = _make_docx_with_header_logo(descr="普通图片")
        result = replace_tagged_images(
            content, {"bid.image:company.logo": LOGO_BYTES}
        )
        with zipfile.ZipFile(BytesIO(result)) as zf:
            assert zf.read(media_part) == PNG_1X1


# ---------- 招标方变量 ----------

@pytest.mark.django_db
class TestTendererVariables:
    def test_lot_fields_in_context(self, monkeypatch):
        store = _patch_storage(monkeypatch)
        user = User.objects.create_user(username="tv-user", password="pass")
        project = Project.objects.create(name="招标方项目", created_by=user)
        lot = Lot.objects.create(
            name="标段", project=project,
            tenderer="某某银行股份有限公司",
            agent="某某招标代理",
            bid_deadline="2026-09-01 09:30",
        )
        outline = Outline.objects.create(
            project=project, lot=lot, name="招标方标书",
            source="preset", created_by=user,
        )
        Section.objects.create(
            outline=outline, title="章节", level=1, sort_order=0, content="内容",
        )

        content = _make_docx([
            ("bid.var:project.tenderer", "招标人", "招标人"),
            ("bid.var:project.agent", "代理机构", "代理机构"),
            ("bid.var:project.bid_deadline", "截止时间", "截止时间"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ])
        template = BidWordTemplate.objects.create(
            name="招标方模板", code="tv-tpl", scope_type="system",
        )
        version = BidWordTemplateVersion.objects.create(
            template=template, version_no=1,
            object_key=f"bid-templates/system/{template.id}/versions/v1.docx",
            file_name="v1.docx", file_size=len(content),
            file_hash=f"tv-hash-{template.id}",
        )
        store[version.object_key] = content

        document, _ = render_bid_document(
            template=template, version=version, outline=outline, user=user,
        )
        text = _docx_text(store[document.object_key])
        assert "某某银行股份有限公司" in text
        assert "某某招标代理" in text
        assert "2026-09-01 09:30" in text
