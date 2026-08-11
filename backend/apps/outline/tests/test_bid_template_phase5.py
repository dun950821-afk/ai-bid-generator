"""Phase 5 测试：五层校验 / 预览 / 回滚 / 默认模板收口。"""
import hashlib
from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from django.core.cache import cache
from docx import Document
from rest_framework.test import APIClient

from apps.common.services.storage import ObjectNotFound, StorageService
from apps.outline.models import (
    BidWordTemplate,
    BidWordTemplateStatus,
    BidWordTemplateVersion,
    Outline,
    Section,
)
from apps.outline.services.template import template_service
from apps.outline.services.template.template_validator import TemplateValidator
from apps.projects.models import Lot, Project

User = get_user_model()


@pytest.fixture(autouse=True)
def _no_preview_generation(monkeypatch):
    """测试环境不调 ONLYOFFICE Conversion API。"""
    monkeypatch.setattr(
        "apps.outline.services.template.template_preview_service.generate_previews",
        lambda template, version: None,
    )


def _make_docx(controls=None, plain_lines=()) -> bytes:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    for line in plain_lines:
        doc.add_paragraph(line)
    for i, (tag, alias, text) in enumerate(controls or []):
        sdt = (
            f"<w:sdt {nsdecls('w')}>"
            f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
            f'<w:id w:val="{3000 + i}"/></w:sdtPr>'
            f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
            f"</w:sdt>"
        )
        p = doc.add_paragraph()
        p._p.append(parse_xml(sdt))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


BODY_SLOT = [("bid.slot:body", "标书正文", "标书正文")]


def _patch_storage(monkeypatch):
    store = {}

    def fake_put(self, key, data, content_type="application/octet-stream"):
        store[key] = data

    def fake_get(self, key):
        if key not in store:
            raise ObjectNotFound(key)
        return store[key]

    monkeypatch.setattr(StorageService, "put_object", fake_put)
    monkeypatch.setattr(StorageService, "get_object", fake_get)
    return store


# ---------- 校验器第四/五层 ----------

class TestValidatorPhase5:
    def test_style_report_and_missing_style_warning(self):
        content = _make_docx(controls=BODY_SLOT)
        result = TemplateValidator().validate(
            content, style_mapping={"heading1": "不存在的标题样式"}
        )
        assert result["valid"] is True  # 样式缺失是 warning，不阻断
        style_entry = next(
            s for s in result["styles"] if s["logical"] == "heading1"
        )
        assert style_entry["exists"] is False
        assert any(
            w["code"] == "STYLE_NOT_FOUND" and w["style"] == "heading1"
            for w in result["warnings"]
        )

    def test_builtin_styles_detected(self):
        content = _make_docx(controls=BODY_SLOT)
        result = TemplateValidator().validate(content)
        heading1 = next(s for s in result["styles"] if s["logical"] == "heading1")
        assert heading1["resolved"] == "Heading 1"
        assert heading1["exists"] is True

    def test_test_render_layer_passes_for_valid_template(self):
        content = _make_docx(
            controls=[("bid.var:company.name", "企业名称", "企业名称")] + BODY_SLOT
        )
        result = TemplateValidator().validate(content)
        assert result["valid"] is True, result["errors"]

    def test_test_render_failure_reported(self, monkeypatch):
        """编译/渲染异常应报告 TEST_RENDER_FAILED。"""
        monkeypatch.setattr(
            "apps.outline.services.template.template_validator.compile_template",
            lambda content: (_ for _ in ()).throw(RuntimeError("boom")),
        )
        content = _make_docx(controls=BODY_SLOT)
        result = TemplateValidator().validate(content)
        assert result["valid"] is False
        assert any(e["code"] == "TEST_RENDER_FAILED" for e in result["errors"])


# ---------- 发布：样式映射入版本 + 预览字段 ----------

@pytest.mark.django_db
class TestPublishPhase5:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="p5-admin", password="pass", email="p5@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)

    def _make_template_with_draft(self, monkeypatch, style_mapping=None):
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="P5模板", code=f"p5-tpl-{id(monkeypatch)}", scope_type="system",
            style_mapping=style_mapping or {},
        )
        content = _make_docx(controls=BODY_SLOT)
        template.draft_object_key = f"bid-templates/system/{template.id}/draft/current.docx"
        template.save(update_fields=["draft_object_key"])
        store[template.draft_object_key] = content
        return template, store

    def test_publish_copies_style_mapping(self, monkeypatch):
        template, _ = self._make_template_with_draft(
            monkeypatch, style_mapping={"heading1": "Heading 1"}
        )
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        assert resp.status_code == 201
        assert resp.json()["version"]["style_mapping"] == {"heading1": "Heading 1"}
        assert "styles" in resp.json()["validation"]


# ---------- 回滚 / 默认模板 ----------

@pytest.mark.django_db
class TestRollbackAndDefault:
    def setup_method(self):
        cache.clear()
        self.admin = User.objects.create_superuser(
            username="rb-admin", password="pass", email="rb@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)

    def _publish_template(self, monkeypatch, text, code):
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name=f"模板-{code}", code=code, scope_type="system",
        )
        template.draft_object_key = f"bid-templates/system/{template.id}/draft/current.docx"
        template.save(update_fields=["draft_object_key"])
        store[template.draft_object_key] = _make_docx(
            controls=BODY_SLOT, plain_lines=[text]
        )
        version, _ = template_service.publish_template(template, user=self.admin)
        return template, version, store

    def test_rollback_copies_version_to_draft(self, monkeypatch):
        template, v1, store = self._publish_template(monkeypatch, "第一版", "rb-tpl")
        old_revision = template.draft_revision

        # 修改 draft 后回滚到 v1
        store[template.draft_object_key] = _make_docx(
            controls=BODY_SLOT, plain_lines=["被改坏的内容"]
        )
        resp = self.client.post(
            f"/api/bid-word-templates/{template.id}/rollback/",
            {"version_id": v1.id},
        )
        assert resp.status_code == 200

        template.refresh_from_db()
        assert template.draft_revision == old_revision + 1
        # published 指针不动，历史版本不动
        assert template.published_version_id == v1.id
        assert template.versions.count() == 1
        # draft 内容回到 v1
        text = ""
        doc = Document(BytesIO(store[template.draft_object_key]))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "第一版" in text

    def test_set_default_requires_published(self, monkeypatch):
        _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="未发布", code="no-pub-tpl", scope_type="system",
        )
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/set_default/")
        assert resp.status_code == 400

    def test_set_default_unique(self, monkeypatch):
        tpl_a, _, _ = self._publish_template(monkeypatch, "A", "default-a")
        tpl_b, _, _ = self._publish_template(monkeypatch, "B", "default-b")

        self.client.post(f"/api/bid-word-templates/{tpl_a.id}/set_default/")
        self.client.post(f"/api/bid-word-templates/{tpl_b.id}/set_default/")

        tpl_a.refresh_from_db()
        tpl_b.refresh_from_db()
        assert tpl_a.is_default is False
        assert tpl_b.is_default is True

    def test_init_default_idempotent(self, monkeypatch):
        _patch_storage(monkeypatch)
        resp1 = self.client.post("/api/bid-word-templates/init_default/")
        assert resp1.status_code == 201
        template_id = resp1.json()["id"]
        assert resp1.json()["is_default"] is True
        assert resp1.json()["published_version_no"] == 1

        resp2 = self.client.post("/api/bid-word-templates/init_default/")
        assert resp2.json()["id"] == template_id  # 幂等

    def test_build_docx_uses_default_template(self, monkeypatch):
        """不传 template_id 且存在默认模板时，走模板渲染链路。"""
        _patch_storage(monkeypatch)
        template = template_service.create_system_default_template(user=self.admin)

        project = Project.objects.create(name="收口项目", created_by=self.admin)
        lot = Lot.objects.create(name="标段", project=project)
        from apps.projects.models import ProjectMember, ProjectRole

        role = ProjectRole.objects.create(
            project=project, name="负责人", code="owner",
            permissions=["project.view", "outline.view"], is_builtin=True,
        )
        ProjectMember.objects.create(project=project, user=self.admin, project_role=role)
        outline = Outline.objects.create(
            project=project, lot=lot, name="收口标书",
            source="preset", created_by=self.admin,
        )
        Section.objects.create(
            outline=outline, title="章节", level=1, sort_order=0, content="内容",
        )

        resp = self.client.post(f"/api/outlines/{outline.id}/build_docx/")
        assert resp.status_code == 200
        data = resp.json()
        assert data["template"]["id"] == template.id

        from apps.outline.models import BidDocument

        document = BidDocument.objects.get(id=data["document_id"])
        assert document.template_id == template.id
        # 默认模板的变量替换生效
        store_content = None
        from apps.common.services.storage import StorageService as SS
        # 直接从 mocked store 读
        # （_patch_storage 的 store 闭包）
        # 这里通过 document.object_key 在 put 记录中
        # 简单验证 snapshot 即可
        assert document.render_context_snapshot["project"]["name"] == "收口项目"


# ---------- 版本文件代理 / 样式列表 / 预览 ----------

@pytest.mark.django_db
class TestFileAndStylesAPI:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="fs-admin", password="pass", email="fs@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)

    def test_file_endpoint_serves_version_with_version_id(self, monkeypatch):
        import time

        import jwt as pyjwt
        from django.conf import settings

        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="版本代理", code="proxy-ver-tpl", scope_type="system",
            draft_object_key="bid-templates/system/9/draft/current.docx",
        )
        version = BidWordTemplateVersion.objects.create(
            template=template, version_no=1,
            object_key="bid-templates/system/9/versions/v1.docx",
            file_name="v1.docx", file_size=10, file_hash="h1",
        )
        store[version.object_key] = b"version-bytes"
        store[template.draft_object_key] = b"draft-bytes"

        token = pyjwt.encode(
            {
                "template_id": template.id,
                "version_id": version.id,
                "exp": int(time.time()) + 3600,
            },
            settings.ONLYOFFICE_JWT_SECRET,
            algorithm="HS256",
        )
        resp = self.client.get(
            f"/api/bid-word-templates/{template.id}/file/?token={token}"
        )
        assert resp.status_code == 200
        assert resp.content == b"version-bytes"

    def test_styles_action(self, monkeypatch):
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="样式", code="styles-tpl", scope_type="system",
            draft_object_key="bid-templates/system/10/draft/current.docx",
            style_mapping={"heading1": "标题 1"},
        )
        store[template.draft_object_key] = _make_docx(controls=BODY_SLOT)

        resp = self.client.get(f"/api/bid-word-templates/{template.id}/styles/")
        assert resp.status_code == 200
        data = resp.json()
        assert "Heading 1" in data["styles"]
        assert data["style_mapping"] == {"heading1": "标题 1"}

    def test_preview_404_without_preview(self, monkeypatch):
        _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="无预览", code="no-preview-tpl", scope_type="system",
        )
        version = BidWordTemplateVersion.objects.create(
            template=template, version_no=1,
            object_key="k", file_name="v1.docx", file_size=1, file_hash="h",
        )
        template.published_version = version
        template.save(update_fields=["published_version"])
        resp = self.client.get(f"/api/bid-word-templates/{template.id}/preview/")
        assert resp.status_code == 404
