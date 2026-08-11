"""模板编译器 / 扫描器 / 校验器测试（Phase 2）。"""
import zipfile
from io import BytesIO

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from apps.outline.services.template.template_compiler import (
    compile_template,
    scan_template,
)
from apps.outline.services.template.template_validator import TemplateValidator
from apps.outline.services.template.template_variable_registry import registry


def _sdt_xml(tag: str, alias: str, text: str, sdt_id: int = 1000) -> str:
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/><w:lock w:val="sdtContentLocked"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def make_template(controls, plain_text=None) -> bytes:
    """构造带内容控件的测试模板。

    controls: [(tag, alias, text)]，每个控件独占一个段落。
    """
    doc = Document()
    for i, (tag, alias, text) in enumerate(controls):
        p = doc.add_paragraph()
        p._p.append(parse_xml(_sdt_xml(tag, alias, text, sdt_id=1000 + i)))
    if plain_text:
        doc.add_paragraph(plain_text)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _document_text(content: bytes) -> str:
    """提取 docx 主文档的全部文本。"""
    with zipfile.ZipFile(BytesIO(content)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return xml


BODY_SLOT = ("bid.slot:body", "标书正文", "标书正文")
COMPANY_NAME = ("bid.var:company.name", "企业名称", "企业名称")


# ---------- 扫描器 ----------

class TestScanTemplate:
    def test_scan_var_controls(self):
        content = make_template([COMPANY_NAME, BODY_SLOT])
        result = scan_template(content)
        keys = {(c["type"], c["key"]) for c in result["controls"]}
        assert ("var", "company.name") in keys
        assert ("slot", "body") in keys
        assert result["body_slot_count"] == 1

    def test_scan_raw_jinja_variables(self):
        content = make_template([BODY_SLOT], plain_text="{{ company.name }}")
        result = scan_template(content)
        assert "company.name" in result["raw_variables"]

    def test_scan_no_body_slot(self):
        content = make_template([COMPANY_NAME])
        result = scan_template(content)
        assert result["body_slot_count"] == 0


# ---------- 编译器 ----------

class TestCompileTemplate:
    def test_var_control_compiled_to_jinja(self):
        content = make_template([COMPANY_NAME, BODY_SLOT])
        compiled = compile_template(content)
        text = _document_text(compiled)
        assert "{{ company.name }}" in text
        assert "bid.var:company.name" not in text

    def test_body_slot_compiled_to_paragraph_directive(self):
        content = make_template([COMPANY_NAME, BODY_SLOT])
        compiled = compile_template(content)
        text = _document_text(compiled)
        assert "{{p body }}" in text
        assert "bid.slot:body" not in text

    def test_material_and_image_controls(self):
        content = make_template([
            ("bid.material:business_license", "营业执照", "营业执照"),
            ("bid.image:company.logo", "企业 Logo", "Logo"),
            BODY_SLOT,
        ])
        compiled = compile_template(content)
        text = _document_text(compiled)
        assert "{{ materials.business_license }}" in text
        assert "{{ images.company_logo }}" in text

    def test_compiled_template_renderable_by_docxtpl(self):
        """编译产物必须能被 docxtpl 正常渲染（冒烟测试）。"""
        from docxtpl import DocxTemplate

        content = make_template([COMPANY_NAME])
        compiled = compile_template(content)

        tpl = DocxTemplate(BytesIO(compiled))
        tpl.render({"company": {"name": "测试公司"}}, autoescape=True)
        buffer = BytesIO()
        tpl.save(buffer)
        text = _document_text(buffer.getvalue())
        assert "测试公司" in text

    def test_compile_preserves_other_parts(self):
        """编译不能破坏文档其余结构（styles/media 等 part 保留）。"""
        content = make_template([COMPANY_NAME])
        compiled = compile_template(content)
        with zipfile.ZipFile(BytesIO(compiled)) as zf:
            names = zf.namelist()
        assert "word/styles.xml" in names
        # 编译产物仍是合法 docx
        Document(BytesIO(compiled))


# ---------- 校验器 ----------

class TestTemplateValidator:
    def setup_method(self):
        self.validator = TemplateValidator()

    def test_valid_template_passes(self):
        content = make_template([COMPANY_NAME, BODY_SLOT])
        result = self.validator.validate(content)
        assert result["valid"] is True
        assert result["errors"] == []
        assert "company.name" in result["variables"]

    def test_missing_body_slot_rejected(self):
        content = make_template([COMPANY_NAME])
        result = self.validator.validate(content)
        assert result["valid"] is False
        codes = {e["code"] for e in result["errors"]}
        assert "BODY_SLOT_MISSING" in codes

    def test_duplicate_body_slot_rejected(self):
        content = make_template([BODY_SLOT, BODY_SLOT])
        result = self.validator.validate(content)
        assert result["valid"] is False
        codes = {e["code"] for e in result["errors"]}
        assert "BODY_SLOT_DUPLICATED" in codes

    def test_unknown_variable_rejected(self):
        content = make_template([
            ("bid.var:company.not_exist", "不存在", "不存在"),
            BODY_SLOT,
        ])
        result = self.validator.validate(content)
        assert result["valid"] is False
        assert any(
            e["code"] == "VARIABLE_UNKNOWN" and "company.not_exist" in e["message"]
            for e in result["errors"]
        )

    def test_unknown_raw_jinja_variable_rejected(self):
        content = make_template([BODY_SLOT], plain_text="{{ hacker.invalid }}")
        result = self.validator.validate(content)
        assert result["valid"] is False
        assert any(e["code"] == "VARIABLE_UNKNOWN" for e in result["errors"])

    def test_invalid_docx_rejected(self):
        result = self.validator.validate(b"not a zip at all")
        assert result["valid"] is False
        assert result["errors"][0]["code"] == "DOCX_INVALID"

    def test_material_control_accepted(self):
        content = make_template([
            ("bid.material:business_license", "营业执照", "营业执照"),
            BODY_SLOT,
        ])
        result = self.validator.validate(content)
        assert result["valid"] is True
        assert "material:business_license" in result["variables"]


# ---------- 变量注册中心 ----------

class TestTemplateVariableRegistry:
    def test_keys_unique(self):
        keys = [v.key for v in registry.all()]
        assert len(keys) == len(set(keys))

    def test_company_variables_present(self):
        assert registry.get("company.name") is not None
        assert registry.get("company.credit_code") is not None

    def test_control_tag_format(self):
        definition = registry.get("company.name")
        assert definition.control_tag == "bid.var:company.name"
        body = registry.get("body")
        assert body.control_tag == "bid.slot:body"

    def test_parse_control_tag(self):
        assert registry.parse_control_tag("bid.var:project.name") == {
            "type": "var",
            "key": "project.name",
        }
        assert registry.parse_control_tag("bid.material:license") == {
            "type": "material",
            "key": "license",
        }
        assert registry.parse_control_tag("random-tag") is None
        assert registry.parse_control_tag("") is None

    def test_is_known_tag(self):
        assert registry.is_known_tag("bid.var:company.name")
        assert registry.is_known_tag("bid.material:anything")
        assert not registry.is_known_tag("bid.var:company.unknown")
        assert not registry.is_known_tag("other:x.y")
