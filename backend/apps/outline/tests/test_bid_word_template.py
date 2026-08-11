"""Word 模板中心 Phase 1 测试：CRUD / 上传 / 发布 / 下载 / 权限 / 可见性。"""
import hashlib
from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from rest_framework.test import APIClient

from apps.accounts.models import Permission, Role
from apps.common.services.storage import ObjectNotFound, StorageService
from apps.outline.models import (
    BidWordTemplate,
    BidWordTemplateStatus,
)
from apps.projects.models import Lot, Project, ProjectMember, ProjectRole

User = get_user_model()


@pytest.fixture(autouse=True)
def _no_preview_generation(monkeypatch):
    """测试环境不调 ONLYOFFICE Conversion API。"""
    monkeypatch.setattr(
        "apps.outline.services.template.template_preview_service.generate_previews",
        lambda template, version: None,
    )


def _make_docx_bytes(text="模板内容", with_body_slot=True) -> bytes:
    """生成一个真实的 docx 文件字节。

    with_body_slot：追加 bid.slot:body 内容控件（发布校验要求模板
    必须包含一个正文插槽）。
    """
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    doc.add_paragraph(text)
    if with_body_slot:
        sdt = (
            f"<w:sdt {nsdecls('w')}>"
            f'<w:sdtPr><w:alias w:val="标书正文"/><w:tag w:val="bid.slot:body"/>'
            f'<w:id w:val="9999"/></w:sdtPr>'
            f'<w:sdtContent><w:r><w:t>标书正文</w:t></w:r></w:sdtContent>'
            f"</w:sdt>"
        )
        p = doc.add_paragraph()
        p._p.append(parse_xml(sdt))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _grant_view_permission(user):
    perm = Permission.objects.get(code="bid_template.view")
    role = Role.objects.create(code=f"tpl-viewer-{user.username}", name="模板查看")
    role.permissions.add(perm)
    user.roles.add(role)


def _patch_storage(monkeypatch):
    """mock MinIO：记录 put 内容，get 返回最近一次 put 的内容。"""
    store = {}

    def fake_put(self, key, data, content_type="application/octet-stream"):
        store[key] = data

    def fake_get(self, key):
        if key not in store:
            raise ObjectNotFound(key)
        return store[key]

    monkeypatch.setattr(StorageService, "put_object", fake_put)
    monkeypatch.setattr(StorageService, "get_object", fake_get)
    return store


@pytest.mark.django_db
class TestBidWordTemplateCRUD:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="tpl-admin", password="pass", email="a@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)

    def test_create_blank_template(self, monkeypatch):
        _patch_storage(monkeypatch)
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "银行标准模板", "scope_type": "system"},
            format="multipart",
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["code"].startswith("tpl-")
        assert data["has_draft_file"] is True
        assert data["draft_revision"] == 1
        assert data["status"] == BidWordTemplateStatus.DRAFT

        template = BidWordTemplate.objects.get(id=data["id"])
        assert template.draft_object_key.startswith(
            f"bid-templates/system/{template.id}/draft/"
        )

    def test_create_with_upload(self, monkeypatch):
        _patch_storage(monkeypatch)
        from django.core.files.uploadedfile import SimpleUploadedFile

        upload = SimpleUploadedFile("tpl.docx", _make_docx_bytes())
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "上传模板", "scope_type": "system", "file": upload},
            format="multipart",
        )
        assert resp.status_code == 201
        assert resp.json()["has_draft_file"] is True

    def test_create_with_invalid_file_returns_400_and_rolls_back(self, monkeypatch):
        _patch_storage(monkeypatch)
        from django.core.files.uploadedfile import SimpleUploadedFile

        upload = SimpleUploadedFile("evil.txt", b"not a docx")
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "坏模板", "scope_type": "system", "file": upload},
            format="multipart",
        )
        assert resp.status_code == 400
        assert BidWordTemplate.objects.filter(name="坏模板").count() == 0

    def test_create_enterprise_template_requires_enterprise(self):
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "企业模板", "scope_type": "enterprise"},
            format="multipart",
        )
        assert resp.status_code == 400

    def test_view_only_user_cannot_create(self):
        viewer = User.objects.create_user(username="tpl-viewer", password="pass")
        _grant_view_permission(viewer)
        client = APIClient()
        client.force_authenticate(user=viewer)
        resp = client.post(
            "/api/bid-word-templates/",
            {"name": "x", "scope_type": "system"},
            format="multipart",
        )
        assert resp.status_code == 403

    def test_unauthenticated_rejected(self):
        resp = APIClient().get("/api/bid-word-templates/")
        assert resp.status_code in (401, 403)


@pytest.mark.django_db
class TestBidWordTemplateVisibility:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="vis-admin", password="pass", email="v@a.com"
        )
        self.owner = User.objects.create_user(username="vis-owner", password="pass")
        self.outsider = User.objects.create_user(username="vis-outsider", password="pass")
        _grant_view_permission(self.owner)
        _grant_view_permission(self.outsider)

        project = Project.objects.create(name="p1", created_by=self.owner)
        lot = Lot.objects.create(name="l1", project=project)
        role = ProjectRole.objects.create(
            project=project, name="负责人", code="owner",
            permissions=["project.view"], is_builtin=True,
        )
        ProjectMember.objects.create(
            project=project, user=self.owner, project_role=role
        )

        self.system_tpl = BidWordTemplate.objects.create(
            name="系统模板", code="sys-tpl", scope_type="system"
        )
        self.project_tpl = BidWordTemplate.objects.create(
            name="项目模板", code="proj-tpl", scope_type="project", project=project
        )

    def _list_names(self, user):
        client = APIClient()
        client.force_authenticate(user=user)
        resp = client.get("/api/bid-word-templates/")
        assert resp.status_code == 200
        payload = resp.json()
        items = payload.get("results", payload) if isinstance(payload, dict) else payload
        return {item["name"] for item in items}

    def test_member_sees_system_and_project_templates(self):
        names = self._list_names(self.owner)
        assert "系统模板" in names
        assert "项目模板" in names

    def test_outsider_cannot_see_project_template(self):
        names = self._list_names(self.outsider)
        assert "系统模板" in names
        assert "项目模板" not in names


@pytest.mark.django_db
class TestBidWordTemplatePublishAndDownload:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="pub-admin", password="pass", email="p@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)
        self.docx_bytes = _make_docx_bytes()

    def _create_template_with_file(self, monkeypatch):
        store = _patch_storage(monkeypatch)
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "发布模板", "scope_type": "system"},
            format="multipart",
        )
        assert resp.status_code == 201
        template = BidWordTemplate.objects.get(id=resp.json()["id"])
        # 模拟 OnlyOffice 保存后的内容
        store[template.draft_object_key] = self.docx_bytes
        return template, store

    def test_publish_creates_immutable_version(self, monkeypatch):
        template, store = self._create_template_with_file(monkeypatch)
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        assert resp.status_code == 201
        data = resp.json()
        assert data["version"]["version_no"] == 1
        assert data["version"]["file_hash"] == hashlib.sha256(self.docx_bytes).hexdigest()
        assert data["validation"]["valid"] is True

        template.refresh_from_db()
        assert template.status == BidWordTemplateStatus.ACTIVE
        assert template.published_version.version_no == 1
        # 版本文件已写入版本路径
        assert template.published_version.object_key.endswith("/versions/v1.docx")
        assert store[template.published_version.object_key] == self.docx_bytes

    def test_publish_increments_version_no(self, monkeypatch):
        template, _ = self._create_template_with_file(monkeypatch)
        self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        assert resp.status_code == 201
        assert resp.json()["version"]["version_no"] == 2

    def test_publish_without_file_returns_400(self, monkeypatch):
        _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="空模板", code="empty-tpl", scope_type="system"
        )
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        assert resp.status_code == 400

    def test_publish_rejects_template_without_body_slot(self, monkeypatch):
        """无正文插槽的模板不允许发布，返回完整校验结果。"""
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="无插槽", code="no-slot-tpl", scope_type="system",
            draft_object_key="bid-templates/system/x/draft/current.docx",
        )
        store[template.draft_object_key] = _make_docx_bytes(with_body_slot=False)
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        assert resp.status_code == 400
        data = resp.json()
        assert data["validation"]["valid"] is False
        codes = {e["code"] for e in data["validation"]["errors"]}
        assert "BODY_SLOT_MISSING" in codes
        assert template.versions.count() == 0

    def test_validate_action_returns_result(self, monkeypatch):
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="校验模板", code="validate-tpl", scope_type="system",
            draft_object_key="bid-templates/system/y/draft/current.docx",
        )
        store[template.draft_object_key] = _make_docx_bytes()
        resp = self.client.post(f"/api/bid-word-templates/{template.id}/validate/")
        assert resp.status_code == 200
        assert resp.json()["valid"] is True

    def test_variables_action_scans_draft(self, monkeypatch):
        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="扫描模板", code="scan-tpl", scope_type="system",
            draft_object_key="bid-templates/system/z/draft/current.docx",
        )
        store[template.draft_object_key] = _make_docx_bytes()
        resp = self.client.get(f"/api/bid-word-templates/{template.id}/variables/")
        assert resp.status_code == 200
        data = resp.json()
        assert data["body_slot_count"] == 1

    def test_variable_registry_api(self):
        resp = self.client.get("/api/bid-word-template-variables/")
        assert resp.status_code == 200
        groups = resp.json()["groups"]
        categories = {g["category"] for g in groups}
        assert {"project", "company", "system", "document", "special"} <= categories
        company = next(g for g in groups if g["category"] == "company")
        name_var = next(
            v for v in company["variables"] if v["key"] == "company.name"
        )
        assert name_var["control_tag"] == "bid.var:company.name"

    def test_versions_list(self, monkeypatch):
        template, _ = self._create_template_with_file(monkeypatch)
        self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        resp = self.client.get(f"/api/bid-word-templates/{template.id}/versions/")
        assert resp.status_code == 200
        payload = resp.json()
        assert len(payload) == 1
        assert payload[0]["version_no"] == 1

    def test_download_draft(self, monkeypatch):
        template, _ = self._create_template_with_file(monkeypatch)
        resp = self.client.get(f"/api/bid-word-templates/{template.id}/download/")
        assert resp.status_code == 200
        assert resp.content == self.docx_bytes

    def test_download_version(self, monkeypatch):
        template, _ = self._create_template_with_file(monkeypatch)
        self.client.post(f"/api/bid-word-templates/{template.id}/publish/")
        template.refresh_from_db()
        resp = self.client.get(
            f"/api/bid-word-templates/{template.id}/download/"
            f"?version_id={template.published_version.id}"
        )
        assert resp.status_code == 200
        assert resp.content == self.docx_bytes


@pytest.mark.django_db
class TestBidWordTemplateEditor:
    def setup_method(self):
        self.admin = User.objects.create_superuser(
            username="ed-admin", password="pass", email="e@a.com"
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.admin)

    def test_editor_config_returns_callback_and_file_url(self, monkeypatch):
        _patch_storage(monkeypatch)
        resp = self.client.post(
            "/api/bid-word-templates/",
            {"name": "编辑模板", "scope_type": "system"},
            format="multipart",
        )
        template_id = resp.json()["id"]

        resp = self.client.get(
            f"/api/bid-word-templates/{template_id}/editor_config/"
        )
        assert resp.status_code == 200
        data = resp.json()
        config = data["config"]
        assert data["documentServerUrl"]
        assert config["document"]["key"]
        assert f"/api/bid-word-templates/{template_id}/file/" in config["document"]["url"]
        assert (
            f"/api/onlyoffice/callback/template/{template_id}/"
            in config["editorConfig"]["callbackUrl"]
        )
        assert config["token"]

    def test_editor_config_without_file_returns_400(self):
        template = BidWordTemplate.objects.create(
            name="无文件", code="nofile-tpl", scope_type="system"
        )
        resp = self.client.get(
            f"/api/bid-word-templates/{template.id}/editor_config/"
        )
        assert resp.status_code == 400

    def test_file_endpoint_requires_valid_jwt(self, monkeypatch):
        import time

        import jwt as pyjwt
        from django.conf import settings

        store = _patch_storage(monkeypatch)
        template = BidWordTemplate.objects.create(
            name="代理模板", code="proxy-tpl", scope_type="system",
            draft_object_key="bid-templates/system/1/draft/current.docx",
        )
        store[template.draft_object_key] = self.docx_bytes = _make_docx_bytes()

        # 无 token → 403
        resp = self.client.get(f"/api/bid-word-templates/{template.id}/file/")
        assert resp.status_code == 403

        # 合法 token → 200
        token = pyjwt.encode(
            {"template_id": template.id, "exp": int(time.time()) + 3600},
            settings.ONLYOFFICE_JWT_SECRET,
            algorithm="HS256",
        )
        resp = self.client.get(
            f"/api/bid-word-templates/{template.id}/file/?token={token}"
        )
        assert resp.status_code == 200
        assert resp.content == self.docx_bytes
