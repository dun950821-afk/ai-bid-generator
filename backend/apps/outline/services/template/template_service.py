# backend/apps/outline/services/template/template_service.py
"""Word 模板服务：draft 文件管理与发布。

draft / 版本对象键结构（方案 §53）：
    bid-templates/{scope}/{template_id}/draft/current.docx
    bid-templates/{scope}/{template_id}/versions/v{n}.docx
"""

import hashlib
import logging
from io import BytesIO
from typing import Tuple

from django.utils import timezone

from apps.common.services.file_magic import extension_of, is_allowed_upload
from apps.common.services.storage import StorageService
from apps.outline.models import (
    BidWordTemplate,
    BidWordTemplateScope,
    BidWordTemplateStatus,
    BidWordTemplateVersion,
    BidWordTemplateVersionStatus,
)

logger = logging.getLogger(__name__)

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# 模板文件上限：20MB（标书模板一般只有样式和封面，远大于此即异常）
MAX_TEMPLATE_FILE_SIZE = 20 * 1024 * 1024


class TemplateValidationError(ValueError):
    """模板文件不合法。"""

    def __init__(self, message, validation: dict | None = None):
        super().__init__(message)
        self.validation = validation


def build_draft_object_key(template: BidWordTemplate) -> str:
    return f"bid-templates/{template.scope_type}/{template.id}/draft/current.docx"


def build_version_object_key(template: BidWordTemplate, version_no: int) -> str:
    return (
        f"bid-templates/{template.scope_type}/{template.id}"
        f"/versions/v{version_no}.docx"
    )


def file_sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def validate_docx_upload(filename: str, content: bytes) -> None:
    """校验上传的模板文件是合法 docx。

    Raises:
        TemplateValidationError: 文件类型/大小不合法
    """
    if len(content) > MAX_TEMPLATE_FILE_SIZE:
        raise TemplateValidationError("模板文件超过 20MB 限制")
    if len(content) == 0:
        raise TemplateValidationError("模板文件为空")
    if extension_of(filename) != "docx":
        raise TemplateValidationError("模板文件必须是 .docx 格式")
    if not is_allowed_upload(filename, content[:4096]):
        raise TemplateValidationError(
            "文件内容与扩展名不符或已损坏，请用 Word/WPS 另存为普通 DOCX 后上传"
        )
    # 进一步确认 python-docx 能打开（OOXML 结构完整）
    try:
        from docx import Document

        Document(BytesIO(content))
    except Exception as exc:
        raise TemplateValidationError(f"DOCX 文件结构不合法：{exc}") from exc


def create_blank_draft(template: BidWordTemplate) -> None:
    """为模板创建空白 draft 文件。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    object_key = build_draft_object_key(template)
    StorageService().put_object(object_key, content, content_type=DOCX_CONTENT_TYPE)

    template.draft_object_key = object_key
    template.draft_revision = 1
    template.draft_file_key = template.generate_draft_file_key()
    template.save(
        update_fields=["draft_object_key", "draft_revision", "draft_file_key", "updated_at"]
    )


def save_upload_as_draft(template: BidWordTemplate, filename: str, content: bytes) -> None:
    """把用户上传的 docx 存为模板 draft 文件。

    Raises:
        TemplateValidationError: 文件不合法
    """
    validate_docx_upload(filename, content)

    object_key = build_draft_object_key(template)
    StorageService().put_object(object_key, content, content_type=DOCX_CONTENT_TYPE)

    template.draft_object_key = object_key
    template.draft_revision = template.draft_revision + 1 if template.draft_revision else 1
    template.draft_file_key = template.generate_draft_file_key()
    template.save(
        update_fields=["draft_object_key", "draft_revision", "draft_file_key", "updated_at"]
    )


def save_draft_content(template: BidWordTemplate, content: bytes, bump_revision: bool = True) -> None:
    """ONLYOFFICE 回调保存：更新 draft 文件。

    status=2（编辑完成）：bump_revision=True，修订号 +1 并更换 file_key，
    不产生业务版本（方案 §17/§18）。
    status=6（forcesave）：bump_revision=False，只覆盖文件内容，不换 key。
    """
    if not content:
        raise TemplateValidationError("回调保存的文件内容为空")

    object_key = template.draft_object_key or build_draft_object_key(template)
    StorageService().put_object(object_key, content, content_type=DOCX_CONTENT_TYPE)

    update_fields = ["draft_object_key", "draft_saved_at", "updated_at"]
    template.draft_object_key = object_key
    template.draft_saved_at = timezone.now()
    if bump_revision:
        template.draft_revision += 1
        template.draft_file_key = template.generate_draft_file_key()
        update_fields += ["draft_revision", "draft_file_key"]
    template.save(update_fields=update_fields)


def publish_template(template: BidWordTemplate, user=None) -> Tuple[BidWordTemplateVersion, dict]:
    """发布当前 draft 为新业务版本。

    发布前执行模板校验（Phase 2 接入前三层：docx 完整性 / 变量白名单 /
    正文插槽），校验不通过不允许发布（方案 §50）。

    Returns:
        (version, validation_result)

    Raises:
        TemplateValidationError: 无 draft 文件、文件不合法或校验未通过
            （validation 属性携带完整校验结果）
    """
    from apps.outline.services.template.template_validator import TemplateValidator

    if not template.draft_object_key:
        raise TemplateValidationError("模板还没有文件，请先上传或在线编辑")

    storage = StorageService()
    content = storage.get_object(template.draft_object_key)
    validate_docx_upload(f"{template.code}.docx", content)

    validation_result = TemplateValidator().validate(
        content, style_mapping=template.style_mapping
    )
    if not validation_result["valid"]:
        first_error = validation_result["errors"][0]["message"]
        raise TemplateValidationError(
            f"模板校验未通过：{first_error}", validation=validation_result
        )

    latest = template.versions.order_by("-version_no").first()
    version_no = (latest.version_no + 1) if latest else 1

    version_key = build_version_object_key(template, version_no)
    storage.put_object(version_key, content, content_type=DOCX_CONTENT_TYPE)

    version = BidWordTemplateVersion.objects.create(
        template=template,
        version_no=version_no,
        object_key=version_key,
        file_name=f"{template.name}_v{version_no}.docx",
        file_size=len(content),
        file_hash=file_sha256(content),
        style_mapping=template.style_mapping,
        variable_schema=validation_result["variables"],
        validation_status=BidWordTemplateVersionStatus.PASSED,
        validation_result=validation_result,
        published_at=timezone.now(),
        created_by=user,
    )

    template.published_version = version
    template.status = BidWordTemplateStatus.ACTIVE
    template.save(update_fields=["published_version", "status", "updated_at"])

    logger.info(
        f"Published bid word template: template_id={template.id}, "
        f"version_no={version_no}, hash={version.file_hash[:12]}"
    )

    # 发布时生成预览产物（best-effort：ONLYOFFICE 不可达不阻断发布）
    from apps.outline.services.template.template_preview_service import (
        generate_previews,
    )

    generate_previews(template, version)

    return version, validation_result


def rollback_to_version(template: BidWordTemplate, version_id: int, user=None) -> BidWordTemplateVersion:
    """回滚：把指定历史版本的文件复制为当前 draft（方案 §44）。

    不修改历史版本，也不移动 published 指针——用户检查后可再发布，
    生成新的业务版本。
    """
    version = template.versions.filter(pk=version_id).first()
    if version is None:
        raise TemplateValidationError("目标版本不存在")

    storage = StorageService()
    content = storage.get_object(version.object_key)

    draft_key = build_draft_object_key(template)
    storage.put_object(draft_key, content, content_type=DOCX_CONTENT_TYPE)

    template.draft_object_key = draft_key
    template.draft_revision += 1
    template.draft_file_key = template.generate_draft_file_key()
    template.draft_saved_at = timezone.now()
    template.updated_by = user
    template.save(
        update_fields=[
            "draft_object_key",
            "draft_revision",
            "draft_file_key",
            "draft_saved_at",
            "updated_by",
            "updated_at",
        ]
    )
    logger.info(
        f"Template rollback: template_id={template.id}, "
        f"from version_no={version.version_no}, new_revision={template.draft_revision}"
    )
    return version


def set_default_template(template: BidWordTemplate) -> None:
    """设为默认模板（全局唯一）；必须先有发布版本。"""
    if template.published_version_id is None:
        raise TemplateValidationError("模板还没有已发布的版本，不能设为默认")
    BidWordTemplate.objects.filter(is_default=True).exclude(
        pk=template.pk
    ).update(is_default=False)
    template.is_default = True
    template.save(update_fields=["is_default", "updated_at"])


def get_default_template() -> BidWordTemplate | None:
    """当前默认模板（有发布版本才有效）。"""
    return (
        BidWordTemplate.objects.filter(is_default=True)
        .select_related("published_version")
        .first()
    )


def create_system_default_template(user=None) -> BidWordTemplate:
    """初始化系统默认简洁模板（幂等：已有默认模板则直接返回）。

    用于默认模板收口（方案 §46）：让所有 Word 生成统一走
    TemplateRenderService，消灭裸生成双路径。
    """
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt

    existing = get_default_template()
    if existing and existing.published_version_id:
        return existing

    # 生成默认模板 docx：标题 + 项目/企业变量 + 正文插槽
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("{{ document.title }}")
    run.font.size = Pt(22)
    run.bold = True

    doc.add_paragraph("项目名称：{{ project.name }}")
    doc.add_paragraph("投标人：{{ company.name }}")
    doc.add_paragraph("日期：{{ system.export_date }}")
    doc.add_page_break()

    body_sdt = (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="标书正文"/><w:tag w:val="bid.slot:body"/>'
        f'<w:id w:val="100001"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>标书正文</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )
    p = doc.add_paragraph()
    p._p.append(parse_xml(body_sdt))

    buffer = BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    template = BidWordTemplate.objects.create(
        name="系统默认简洁模板",
        code="system-default",
        description="系统自动生成的默认模板，可在模板中心修改后重新发布",
        scope_type=BidWordTemplateScope.SYSTEM,
        created_by=user,
        updated_by=user,
    )
    save_upload_as_draft(template, "system-default.docx", content)
    publish_template(template, user=user)
    set_default_template(template)
    return template
