# backend/apps/outline/services/document/word_body_renderer.py
"""标书正文渲染器（AST 驱动，方案 §12/§48）。

链路：Markdown → MarkdownParser → AST 节点 → 本渲染器 → Word。

可写入任意 python-docx 文档对象：
- 旧链路：BidDocxBuilder 用空白 Document() 调用本渲染器；
- 模板链路：TemplateRenderService 用 docxtpl subdoc（继承模板样式）调用。

样式契约（方案 §13/§14）：渲染只使用逻辑样式名（heading1~4/body/
list_bullet/list_number/table/quote/image_caption），经 style_mapping
映射到模板实际样式；无映射时回退 Word 内置样式，模板缺样式时
静默降级为无样式，不中断生成。
"""

import logging
from io import BytesIO
from typing import List, Optional

from docx.shared import Mm

from apps.outline.models import Section
from apps.outline.services.document.ast import (
    HeadingNode,
    ImageNode,
    ListNode,
    MaterialNode,
    PageBreakNode,
    ParagraphNode,
    QuoteNode,
    TableNode,
)
from apps.outline.services.document.image_resolver import resolve_image_bytes
from apps.outline.services.document.markdown_parser import MarkdownParser
from apps.outline.services.section_numbering_service import (
    SectionNumberingService,
    strip_number_prefix,
)

logger = logging.getLogger(__name__)

# 逻辑样式 → 内置样式回退
BUILTIN_STYLE_FALLBACK = {
    "list_bullet": "List Bullet",
    "list_number": "List Number",
    "quote": "Intense Quote",
}


class WordBodyRenderer:
    """把章节树内容渲染进指定的文档对象。"""

    def __init__(self, style_mapping: Optional[dict] = None):
        self.style_mapping = style_mapping or {}
        self._doc = None
        self._parser = MarkdownParser()

    def render(
        self,
        doc,
        sections: List[Section],
        material_package=None,
    ) -> None:
        """把全部章节写入 doc。

        Args:
            doc: python-docx Document 或 docxtpl subdoc
            sections: 章节扁平列表
            material_package: 材料包（可选，用于材料占位符插图）
        """
        self._doc = doc
        sections = list(sections)
        number_map = SectionNumberingService().build_number_map(sections)

        # 构建树形结构
        section_map = {}
        root_sections = []
        for section in sections:
            section_map[section.id] = {"section": section, "children": []}
        for section in sections:
            if section.parent_id is None:
                root_sections.append(section_map[section.id])
            elif section.parent_id in section_map:
                section_map[section.parent_id]["children"].append(section_map[section.id])

        def sort_children(nodes):
            nodes.sort(key=lambda x: x["section"].sort_order)
            for node in nodes:
                sort_children(node["children"])

        sort_children(root_sections)

        def write_section(node, depth=0):
            self._add_section(node["section"], depth, number_map, material_package)
            for child in node["children"]:
                write_section(child, depth + 1)

        for node in root_sections:
            write_section(node)

    # ---------- 样式解析 ----------

    def _style_kwargs(self, logical: str, builtin: str = "") -> dict:
        """逻辑样式 → add_paragraph 的 style 参数。

        优先 style_mapping；其次内置样式；模板缺样式时返回空 dict
        （python-docx 对不存在的样式名抛 KeyError，这里提前规避）。
        """
        mapped = self.style_mapping.get(logical)
        if mapped:
            if self._has_style(mapped):
                return {"style": mapped}
        elif builtin and self._has_style(builtin):
            return {"style": builtin}
        return {}

    def _has_style(self, name: str) -> bool:
        try:
            self._doc.styles[name]
            return True
        except (KeyError, AttributeError):
            return False

    def _add_heading(self, text: str, level: int):
        mapped = self.style_mapping.get(f"heading{level}")
        if mapped and self._has_style(mapped):
            return self._doc.add_paragraph(text, style=mapped)
        try:
            return self._doc.add_heading(text, level=level)
        except KeyError:
            # 模板缺少内置 Heading 样式时降级为普通段落
            return self._doc.add_paragraph(text)

    # ---------- 章节 ----------

    def _add_section(self, section: Section, depth: int, number_map: dict, material_package=None):
        heading_level = min(max(depth + 1, 1), 4)
        self._add_heading(self._get_section_title(section, number_map), level=heading_level)

        content = section.content or ""
        if content.strip():
            for node in self._parser.parse(content):
                self._write_node(node, material_package)

    def _get_section_title(self, section: Section, number_map: dict) -> str:
        number = number_map.get(section.id, "")
        title = strip_number_prefix(section.title) if section.title else ""
        if number:
            return f"{number}{title}"
        return section.title or ""

    def _no_indent_chars(self, paragraph) -> None:
        """段落首行缩进清零（firstLineChars=0，插到合法位置）。

        模板的 Normal 样式带首行缩进 2 字符，表格单元格、图片等
        非正文段落必须显式清零，否则会继承缩进。
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        ppr = paragraph._p.get_or_add_pPr()
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            ppr.remove(ind)
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLineChars="0"/>')
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            jc.addprevious(ind)
        else:
            ppr.append(ind)

    # ---------- AST 节点写入 ----------

    def _write_node(self, node, material_package) -> None:
        if isinstance(node, HeadingNode):
            level = min(max(node.level, 1), 4)
            self._add_heading(node.text, level=level)
        elif isinstance(node, ParagraphNode):
            self._write_paragraph(node)
        elif isinstance(node, ListNode):
            self._write_list(node)
        elif isinstance(node, TableNode):
            self._write_table(node)
        elif isinstance(node, ImageNode):
            self._write_image(node)
        elif isinstance(node, MaterialNode):
            self._write_material(node, material_package)
        elif isinstance(node, PageBreakNode):
            self._doc.add_page_break()
        elif isinstance(node, QuoteNode):
            paragraph = self._doc.add_paragraph(
                **self._style_kwargs("quote", BUILTIN_STYLE_FALLBACK["quote"])
            )
            self._write_segments(paragraph, node.segments)

    def _write_paragraph(self, node: ParagraphNode):
        paragraph = self._doc.add_paragraph(**self._style_kwargs("body"))
        self._write_segments(paragraph, node.segments)

    def _write_segments(self, paragraph, segments) -> None:
        for segment in segments:
            if not segment.text:
                continue
            run = paragraph.add_run(segment.text)
            if segment.bold:
                run.bold = True
            if segment.italic:
                run.italic = True

    def _write_list(self, node: ListNode):
        logical = "list_number" if node.ordered else "list_bullet"
        builtin = BUILTIN_STYLE_FALLBACK[logical]
        for item in node.items:
            paragraph = self._doc.add_paragraph(**self._style_kwargs(logical, builtin))
            self._write_segments(paragraph, item)

    def _write_table(self, node: TableNode):
        rows_data = node.rows
        if not rows_data:
            return

        max_cols = max(len(row) for row in rows_data)
        table = self._doc.add_table(rows=len(rows_data), cols=max_cols)
        try:
            table.style = self.style_mapping.get("table", "Table Grid")
        except KeyError:
            pass

        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        # 单元格不继承正文的首行缩进
                        self._no_indent_chars(paragraph)
                        # 表头行加粗；字号交由模板样式控制
                        if node.has_header and i == 0:
                            for run in paragraph.runs:
                                run.bold = True

    def _write_image(self, node: ImageNode):
        """正文图片：URL → MinIO 字节 → 插图；失败降级为占位文字。"""
        image_data = resolve_image_bytes(node.url)
        if image_data is None:
            self._doc.add_paragraph(
                f"【图片无法解析：{node.alt or node.url[:80]}】",
                **self._style_kwargs("body"),
            )
            return

        paragraph = self._doc.add_paragraph()
        paragraph.alignment = 1  # 图片默认居中
        self._no_indent_chars(paragraph)
        run = paragraph.add_run()
        # HTML img 带 width（px）时按 96dpi 换算并封顶 150mm（版心宽度）
        if node.width_px:
            width_mm = min(node.width_px * 25.4 / 96, 150)
        else:
            width_mm = 120
        run.add_picture(BytesIO(image_data), width=Mm(width_mm))

        if node.alt:
            caption = self._doc.add_paragraph(
                node.alt, **self._style_kwargs("image_caption")
            )
            caption.alignment = 1
            self._no_indent_chars(caption)

    def _write_material(self, node: MaterialNode, material_package):
        if not material_package:
            self._doc.add_paragraph(f"【缺少材料：{node.usage_key}】")
            return

        item = material_package.items.filter(
            usage_key=node.usage_key
        ).select_related("material").first()

        if item and item.material and item.material.object_key:
            try:
                self._insert_material_image(item.material, node.is_attachment)
            except Exception as exc:
                logger.warning(f"Material image insert failed: {node.usage_key}, {exc}")
                self._doc.add_paragraph(f"【材料图片插入失败：{node.usage_key}】")
        else:
            self._doc.add_paragraph(f"【缺少材料：{node.usage_key}】")

    def _insert_material_image(self, material, is_attachment: bool = False):
        from apps.common.services.storage import StorageService

        storage = StorageService()
        image_data = storage.get_object(material.object_key)
        image_stream = BytesIO(image_data)

        paragraph = self._doc.add_paragraph()
        if is_attachment:
            paragraph.alignment = 1  # 居中
        self._no_indent_chars(paragraph)

        width_mm = 150 if is_attachment else 100
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Mm(width_mm))

        if is_attachment and material.title:
            caption = self._doc.add_paragraph(
                material.title, **self._style_kwargs("image_caption")
            )
            caption.alignment = 1
            self._no_indent_chars(caption)
