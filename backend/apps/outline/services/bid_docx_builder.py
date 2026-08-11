# backend/apps/outline/services/bid_docx_builder.py
"""标书 Word 文档生成服务（旧链路薄包装）。

过渡期保留：generate_document 不传 template_id 时走这里（从空白
Document() 生成裸文档）。实际的内容渲染逻辑已拆到
services/document/word_body_renderer.py，与模板渲染链路共用；
待 Phase 5 系统默认模板上线后，本模块将整体下线。
"""

from io import BytesIO
from typing import List, Tuple

from django.core.files.base import ContentFile
from docx import Document

from apps.outline.models import Section
from apps.outline.services.document.word_body_renderer import WordBodyRenderer


class BidDocxBuilder:
    """将大纲章节内容组装为 Word docx（无模板裸生成）。"""

    def build(
        self,
        outline,
        sections: List[Section],
    ) -> Tuple[ContentFile, List[dict]]:
        """生成 Word 文档。

        Returns:
            (ContentFile, warnings): Word 文件和警告列表
        """
        warnings = []
        doc = Document()

        # 文档标题
        title_text = outline.name or "投标文件"
        title = doc.add_heading(title_text, level=0)
        title.alignment = 1  # 居中

        # 内容检查
        sections = list(sections)
        sections_with_content = [s for s in sections if s.content and s.content.strip()]
        if not sections_with_content:
            warnings.append({
                "type": "no_content",
                "message": "没有任何章节包含内容，生成的文档将为空",
            })
        else:
            empty_sections = [s for s in sections if not s.content or not s.content.strip()]
            if empty_sections:
                titles = [s.title[:30] for s in empty_sections[:5]]
                warnings.append({
                    "type": "partial_content",
                    "message": f"部分章节内容为空：{', '.join(titles)}{'等' if len(empty_sections) > 5 else ''}",
                })

        # 正文渲染（与模板链路共用同一渲染器）
        WordBodyRenderer().render(doc, sections, self._get_material_package(outline))

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{outline.name or '投标文件'}.docx"
        return ContentFile(buffer.read(), name=filename), warnings

    def _get_material_package(self, outline):
        """获取大纲关联的材料包。"""
        try:
            return outline.material_package
        except Exception:
            return None
