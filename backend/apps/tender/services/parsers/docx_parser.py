# backend/apps/tender/services/parsers/docx_parser.py
"""DOCX 文档解析器。"""

import io
import logging
from apps.tender.services.parsers.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """DOCX 文档解析器。

    使用 python-docx 提取段落、标题、表格和列表，
    转换为 Markdown 格式。
    """

    SUPPORTED_EXTENSIONS = ["docx"]

    def parse(self, content: bytes, filename: str) -> ParseResult:
        """解析 DOCX 文件。

        Args:
            content: DOCX 文件二进制内容
            filename: 文件名

        Returns:
            ParseResult 包含 Markdown 和元数据
        """
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安装，请添加到 requirements.txt")

        document = Document(io.BytesIO(content))

        markdown_parts = []
        table_count = 0
        paragraph_count = 0

        # 遍历文档元素
        for element in document.element.body:
            # 处理段落
            if element.tag.endswith("}p"):
                paragraph = self._find_paragraph_by_element(document, element)
                if paragraph:
                    text = self._process_paragraph(paragraph)
                    if text:
                        markdown_parts.append(text)
                        paragraph_count += 1

            # 处理表格
            elif element.tag.endswith("}tbl"):
                table = self._find_table_by_element(document, element)
                if table:
                    table_md = self._process_table(table)
                    if table_md:
                        markdown_parts.append(table_md)
                        table_count += 1

        markdown = "\n\n".join(markdown_parts)

        # 计算质量指标
        quality_metrics = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "char_count": len(markdown),
            "avg_paragraph_length": len(markdown) // max(paragraph_count, 1),
        }

        # 质量评估
        parse_quality, error_message = self._evaluate_quality(markdown)

        return ParseResult(
            markdown=markdown,
            page_count=0,  # DOCX 无法准确获取页数
            page_map=[],  # DOCX 无法获取页码映射
            parse_engine="docx-python-docx",
            parse_quality=parse_quality,
            quality_metrics=quality_metrics,
            error_message=error_message,
        )

    def _find_paragraph_by_element(self, document, element):
        """根据 XML 元素找到对应的 Paragraph 对象。"""
        from docx.text.paragraph import Paragraph
        for para in document.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_by_element(self, document, element):
        """根据 XML 元素找到对应的 Table 对象。"""
        from docx.table import Table
        for table in document.tables:
            if table._element is element:
                return table
        return None

    def _process_paragraph(self, paragraph) -> str:
        """处理段落，返回 Markdown 文本。"""
        text = paragraph.text.strip()
        if not text:
            return ""

        # 根据样式判断标题级别
        style_name = paragraph.style.name if paragraph.style else ""

        if "Heading 1" in style_name or "标题 1" in style_name:
            return f"# {text}"
        elif "Heading 2" in style_name or "标题 2" in style_name:
            return f"## {text}"
        elif "Heading 3" in style_name or "标题 3" in style_name:
            return f"### {text}"
        elif "Heading 4" in style_name or "标题 4" in style_name:
            return f"#### {text}"
        elif "Heading 5" in style_name or "标题 5" in style_name:
            return f"##### {text}"
        elif "Heading 6" in style_name or "标题 6" in style_name:
            return f"###### {text}"
        elif "List" in style_name or "列表" in style_name:
            # 列表项处理
            return f"- {text}"

        # 检测可能的标题格式（如：一、xxx 或 1. xxx）
        import re
        if re.match(r"^[一二三四五六七八九十]+、", text):
            return f"# {text}"
        if re.match(r"^\d+\.\s+\S", text):
            # 数字编号可能是标题也可能是列表
            if len(text) < 50:
                return f"## {text}"

        return text

    def _process_table(self, table) -> str:
        """处理表格，转换为 Markdown 格式。"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 获取单元格文本，合并多个段落
                cell_text = " ".join(
                    para.text.strip() for para in cell.paragraphs if para.text.strip()
                )
                # 转义管道符
                cell_text = cell_text.replace("|", "\\|")
                cells.append(cell_text)
            rows.append(cells)

        if not rows:
            return ""

        # 构建 Markdown 表格
        lines = []

        # 表头
        if rows:
            header = rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 数据行
        for row in rows[1:]:
            # 确保列数一致
            while len(row) < len(rows[0]):
                row.append("")
            lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

        return "\n".join(lines)

    def _evaluate_quality(self, markdown: str) -> tuple[str, str | None]:
        """评估解析质量。"""
        if len(markdown) < 500:
            return "poor", "解析文本过短，请确认文档内容是否正确"
        if len(markdown) < 1500:
            return "low", "解析文本较短，可能存在提取问题"
        return "high", None
