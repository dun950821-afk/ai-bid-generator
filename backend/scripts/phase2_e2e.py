# -*- coding: utf-8 -*-
"""Phase 2 端到端验证: 案例库 → REPEAT_TABLE 自动填充 + 单独密封文档。

用法(容器内):
    python scripts/phase2_e2e.py
"""

import os
import sys

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

import django

django.setup()

from django.contrib.auth import get_user_model
from apps.enterprise.models import CompanyCase, CompanyProfile
from apps.response_template.models import TenderResponseTemplate

TEMPLATE_ID = 1


def main():
    # 1. 准备案例数据
    user = get_user_model().objects.filter(is_superuser=True).first()
    company = CompanyProfile.objects.filter(is_default=True).first()
    if not company:
        print("✗ 无默认企业, 先在企业资料中心创建公司")
        return

    cases = [
        ("XX农商银行互联网资产安全众测项目", "XX农村商业银行股份有限公司", "2025.06-2025.09", 15.0,
         "对互联网资产开展安全众测, 输出漏洞报告与整改建议"),
        ("XX银行Web应用渗透测试项目", "XX银行", "2024.03-2024.06", 8.5,
         "Web 应用渗透测试、漏洞复现与修复验证"),
        ("XX集团安全评估服务项目", "XX集团有限公司", "2023.10-2024.02", 20.0,
         "全网安全评估、安全巡检与应急响应"),
    ]
    created = 0
    for name, client, period, amount, scope in cases:
        obj, is_new = CompanyCase.objects.get_or_create(
            company=company, project_name=name,
            defaults={
                "client_name": client, "scope": scope, "amount": amount,
                "created_by": user,
            },
        )
        if is_new:
            created += 1
    print(f"[案例库] 准备完成, 新增 {created} 条, 共 {CompanyCase.objects.filter(company=company).count()} 条")

    # 2. 重新生成响应文件(任务体同步执行)
    from apps.response_template.constants import TemplateStatus
    from apps.response_template.tasks import fill_response_template

    template = TenderResponseTemplate.objects.get(pk=TEMPLATE_ID)
    if template.status != TemplateStatus.CONFIRMED:
        template.status = TemplateStatus.CONFIRMED
        template.save(update_fields=["status"])

    result = fill_response_template.run(TEMPLATE_ID)
    print(f"\n[生成] document_ids: {result['document_ids']}")
    print(f"[生成] 警告数: {len(result['warnings'])}")
    for w in result["warnings"][:15]:
        print(f"  ⚠ {w['block_key']}: {w['message'][:70]}")

    # 3. 校验产物
    from apps.common.services.storage import StorageService
    from docx import Document
    import io

    storage = StorageService()
    docs = list(TenderResponseTemplate.objects.get(pk=TEMPLATE_ID).documents.all())
    print(f"\n[产物] 共 {len(docs)} 个:")
    for d in docs:
        print(f"  · kind={d.kind} status={d.status} name={d.file_name} ({d.file_size} bytes)")

    main_doc = next((d for d in docs if d.kind == "main"), None)
    sep_doc = next((d for d in docs if d.kind == "separate"), None)

    if main_doc:
        data = storage.get_object(main_doc.object_key)
        doc = Document(io.BytesIO(data))
        texts = []
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        # 案例表是否填入
        case_hits = [t for t in texts if "众测" in t or "渗透" in t]
        print(f"\n[校验-主文档] 案例表命中案例文字: {len(case_hits)} 处")
        for h in case_hits[:5]:
            print(f"  · {h[:50]}")
        if case_hits:
            print("  ✅ REPEAT_TABLE 案例自动填充生效")
        else:
            print("  ⚠ 未找到案例填充(检查案例表定位)")

    if sep_doc:
        data = storage.get_object(sep_doc.object_key)
        doc = Document(io.BytesIO(data))
        sep_texts = "".join(p.text for p in doc.paragraphs) + "".join(
            c.text for t in doc.tables for row in t.rows for c in row.cells
        )
        has_price = "高危" in sep_texts or "报价" in sep_texts or "价格" in sep_texts
        print(f"[校验-密封文档] 包含报价内容: {has_price}")
        if has_price:
            print("  ✅ 单独密封文档生成生效")

    print("\n✅ Phase 2 端到端验证完成")


if __name__ == "__main__":
    main()
