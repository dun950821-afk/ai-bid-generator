# backend/scripts/create_general_bid_template.py
"""生成并发布「通用标书模板」。

结构（接合当前模板引擎能力）：
  封面：项目名称/标段名称/投标人/日期 变量
  投标人信息页：企业信息变量表格
  目录页：Word TOC 域（打开后更新域刷新页码）
  正文：bid.slot:body 插槽
  附件：营业执照材料控件
  页眉：项目名称变量（封面页不显示）；页脚：居中页码

排版：A4、正文宋体小四 1.5 倍行距、标题黑体（三号/四号/小四）。

用法：cd backend && .venv/bin/python scripts/create_general_bid_template.py
幂等：已存在同 code 模板时重新上传草稿并发布新版本。
"""

import os
import sys
from io import BytesIO

import django

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings.dev")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
django.setup()

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls, qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

TEMPLATE_CODE = "general-bid"
TEMPLATE_NAME = "通用标书模板"


def sdt(tag: str, alias: str, text: str, sdt_id: int) -> str:
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def add_var_paragraph(doc, tag, alias, text, sdt_id, *, align=None, bold=False,
                      size=None, font_eastasia=None, space_after=6):
    """添加一个以内容控件为主体的段落。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
        _no_indent(p)  # 居中/右对齐段落不做首行缩进（否则会偏）
    p.paragraph_format.space_after = Pt(space_after)
    p._p.append(parse_xml(sdt(tag, alias, text, sdt_id)))
    # 控件内 run 的格式：在 sdtContent 的 run 上设置
    for run in p._p.iter(qn("w:r")):
        if run.find(qn("w:t")) is not None:
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
                run.insert(0, rpr)
            if bold:
                rpr.append(parse_xml(f"<w:b {nsdecls('w')}/>"))
            if size:
                rpr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>'))
            if font_eastasia:
                rpr.append(parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_eastasia}"/>'
                ))
    return p


def set_style_font(style, ascii_font, eastasia, size_pt, bold=False, black=True):
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if black:
        from docx.shared import RGBColor

        style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), eastasia)


def set_first_line_chars(style, chars: int):
    """设置样式首行缩进（单位：字符，200 = 2 字符）。"""
    ppr = style.element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = parse_xml(f"<w:ind {nsdecls('w')}/>")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))


def set_run_black_songti(run, size_pt=None, bold=None):
    """直接写入的 run：宋体黑色。"""
    from docx.shared import RGBColor

    run.font.name = "宋体"
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _no_indent(paragraph):
    """段落首行缩进清零（表格单元格等不需要缩进的场景）。"""
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = parse_xml(f"<w:ind {nsdecls('w')}/>")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), "0")


def add_toc(doc):
    """目录页：TOC 域（w:dirty 让 Word/OnlyOffice 打开时提示刷新）。"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(title)
    run = title.add_run("目  录")
    set_run_black_songti(run, size_pt=16, bold=True)

    toc_p = doc.add_paragraph()
    fld = (
        f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>'
        f'<w:r {nsdecls("w")}><w:t>（目录将在此处生成：打开文档后右键此处选「更新域」）</w:t></w:r>'
        f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>'
    )
    for r in fld.split("</w:r>")[:-1]:
        toc_p._p.append(parse_xml(r + "</w:r>"))


def build_template_docx() -> bytes:
    doc = Document()

    # ---- 页面与基础样式 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True  # 封面不要页眉页脚

    # 正文：宋体小四、1.5 倍行距、两端对齐、首行缩进 2 字符
    normal = doc.styles["Normal"]
    set_style_font(normal, "宋体", "宋体", 12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_first_line_chars(normal, 200)

    # 标题：宋体加粗黑色（覆盖内置样式的默认蓝色），不缩进
    set_style_font(doc.styles["Heading 1"], "宋体", "宋体", 16, bold=True)  # 三号
    set_style_font(doc.styles["Heading 2"], "宋体", "宋体", 14, bold=True)  # 四号
    set_style_font(doc.styles["Heading 3"], "宋体", "宋体", 12, bold=True)  # 小四
    set_style_font(doc.styles["Heading 4"], "宋体", "宋体", 12, bold=True)
    for level, (before, after) in {1: (13, 13), 2: (7, 7), 3: (6, 6), 4: (6, 6)}.items():
        h = doc.styles[f"Heading {level}"]
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.font.italic = False  # 内置 Heading 4 默认斜体，中文标书不用
        set_first_line_chars(h, 0)

    # 引用/图注：宋体黑色非斜体（内置 Quote 斜体、Caption 蓝色，均不符合惯例）
    for name, size, bold in (("Quote", 12, False), ("Intense Quote", 12, False),
                             ("Caption", 9, False)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        set_style_font(st, "宋体", "宋体", size, bold=bold)
        st.font.italic = False
    try:
        doc.styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except KeyError:
        pass

    # 列表/页眉页脚：宋体黑色，不缩进，页眉页脚小五
    for name, size in (("List Bullet", 12), ("List Number", 12),
                       ("List Paragraph", 12), ("Header", 9), ("Footer", 9)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        set_style_font(st, "宋体", "宋体", size)
        set_first_line_chars(st, 0)

    # 目录条目样式（TOC 域生成的条目使用）：宋体黑色小四
    from docx.enum.style import WD_STYLE_TYPE

    for level in (1, 2, 3):
        toc_name = f"TOC {level}"
        try:
            toc_style = doc.styles[toc_name]
        except KeyError:
            try:
                toc_style = doc.styles.add_style(toc_name, WD_STYLE_TYPE.PARAGRAPH)
            except ValueError:
                continue
        set_style_font(toc_style, "宋体", "宋体", 12)
        set_first_line_chars(toc_style, 0)

    # ---- 页眉（非封面页）：项目名称 + 下划线（惯例） ----
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _no_indent(header_p)
    header_p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f"</w:pBdr>"
    ))
    header_p._p.append(parse_xml(sdt("bid.var:project.name", "项目名称", "项目名称", 9001)))
    run = header_p.add_run("  投标文件")
    set_run_black_songti(run, size_pt=9)

    # ---- 页脚：居中页码 ----
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(footer_p)
    footer_p._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>'
    ))
    footer_p._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
    ))
    footer_p._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>'
    ))

    sid = 100

    def nid():
        nonlocal sid
        sid += 1
        return sid

    # ---- 封面 ----
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(title)
    run = title.add_run("投 标 文 件")
    set_run_black_songti(run, size_pt=26, bold=True)  # 一号
    for _ in range(3):
        doc.add_paragraph()

    cover_fields = [
        ("bid.var:project.name", "项目名称", "项目名称：某某项目"),
        ("bid.var:project.code", "项目编号", "项目编号：ZB-XXXX-XXX"),
        ("bid.var:project.package_name", "标段名称", "标段名称：标段一"),
        ("bid.var:company.name", "企业名称", "投 标 人：某某科技有限公司"),
        ("bid.var:system.export_date", "生成日期", "日期：2026年8月11日"),
    ]
    for tag, alias, text in cover_fields:
        add_var_paragraph(
            doc, tag, alias, text, nid(),
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=18,
        )
    doc.add_page_break()

    # ---- 投标人信息页 ----
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(h)
    run = h.add_run("投标人信息")
    set_run_black_songti(run, size_pt=16, bold=True)

    company_fields = [
        ("企业名称", "bid.var:company.name"),
        ("统一社会信用代码", "bid.var:company.credit_code"),
        ("法定代表人", "bid.var:company.legal_representative"),
        ("注册资本", "bid.var:company.registered_capital"),
        ("成立日期", "bid.var:company.established_date"),
        ("注册地址", "bid.var:company.address"),
        ("联系电话", "bid.var:company.phone"),
        ("电子邮箱", "bid.var:company.email"),
        ("开户银行", "bid.var:company.bank_name"),
        ("银行账号", "bid.var:company.bank_account"),
    ]
    table = doc.add_table(rows=len(company_fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, tag) in enumerate(company_fields):
        label_cell = table.rows[i].cells[0]
        label_cell.text = label
        for p in label_cell.paragraphs:
            _no_indent(p)
            for r in p.runs:
                set_run_black_songti(r, size_pt=10.5, bold=True)  # 五号
        value_cell = table.rows[i].cells[1]
        cell_p = value_cell.paragraphs[0]
        _no_indent(cell_p)
        cell_p._p.append(
            parse_xml(sdt(tag, label, label, nid()))
        )
    doc.add_page_break()

    # ---- 目录页 ----
    add_toc(doc)
    doc.add_page_break()

    # ---- 正文插槽 ----
    body_p = doc.add_paragraph()
    body_p._p.append(parse_xml(sdt("bid.slot:body", "标书正文", "标书正文", nid())))
    doc.add_page_break()

    # ---- 附件 ----
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(h2)
    run = h2.add_run("附  件")
    set_run_black_songti(run, size_pt=16, bold=True)

    p = doc.add_paragraph("一、营业执照")
    _no_indent(p)
    for r in p.runs:
        set_run_black_songti(r, bold=True)
    mat_p = doc.add_paragraph()
    mat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(mat_p)
    mat_p._p.append(parse_xml(sdt("bid.material:business_license", "营业执照", "营业执照", nid())))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def main():
    from django.contrib.auth import get_user_model

    from apps.outline.models import BidWordTemplate
    from apps.outline.services.template import template_service

    User = get_user_model()
    admin = User.objects.filter(is_superuser=True).first()

    content = build_template_docx()

    template = BidWordTemplate.objects.filter(code=TEMPLATE_CODE).first()
    if template is None:
        template = BidWordTemplate.objects.create(
            name=TEMPLATE_NAME,
            code=TEMPLATE_CODE,
            description=(
                "系统预置通用模板：封面 + 投标人信息 + 目录 + 正文 + 附件，"
                "A4 宋体小四 1.5 倍行距，标题黑体。可直接在编辑器中修改后重新发布。"
            ),
            scope_type="system",
            created_by=admin,
            updated_by=admin,
        )
        print(f"created template id={template.id}")
    else:
        print(f"template exists id={template.id}, re-upload draft")

    template_service.save_upload_as_draft(template, "通用标书模板.docx", content)
    version, validation = template_service.publish_template(template, user=admin)
    print(f"published v{version.version_no}, valid={validation['valid']}")
    print(f"variables: {validation['variables']}")
    print(f"preview: image={bool(version.preview_image_key)}, pdf={bool(version.preview_pdf_key)}")

    template_service.set_default_template(template)
    print("set as default template")


if __name__ == "__main__":
    main()
