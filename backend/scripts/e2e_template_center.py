# backend/scripts/e2e_template_center.py
"""Word 模板中心端到端测试脚本。

对运行中的系统（http://localhost）走完整链路：
登录 → 上传模板(含内容控件) → 校验 → 发布(真实 OnlyOffice 转换预览)
→ 选模板生成标书 → 下载 docx 断言内容 → 导出 PDF。

数据种子（项目/标段/大纲/章节/企业/材料包）通过 ORM 直接写入，
被测功能全部走 HTTP API。

运行：cd backend && .venv/bin/python scripts/e2e_template_center.py
前置：docker compose 服务全启动；admin 账号可用。
"""

import os
import sys
import time
from io import BytesIO

import django

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings.dev")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
django.setup()

import requests  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls  # noqa: E402

BASE = "http://localhost"
ADMIN = {"username": "admin", "password": "admin123"}

PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)

PASSED = []
FAILED = []


def check(name, condition, detail=""):
    if condition:
        PASSED.append(name)
        print(f"  ✓ {name}")
    else:
        FAILED.append((name, detail))
        print(f"  × {name}  {detail}")


def sdt_xml(tag, alias, text, sdt_id):
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def build_template_docx():
    """构造带内容控件的模板：封面变量 + 材料 + 正文插槽。"""
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = 1

    for i, (tag, alias, text) in enumerate([
        ("bid.var:project.name", "项目名称", "项目名称"),
        ("bid.var:company.name", "企业名称", "企业名称"),
        ("bid.var:company.legal_representative", "法定代表人", "法定代表人"),
        ("bid.var:system.export_date", "生成日期", "生成日期"),
    ]):
        p = doc.add_paragraph()
        p._p.append(parse_xml(sdt_xml(tag, alias, text, 5000 + i)))

    p = doc.add_paragraph()
    p._p.append(parse_xml(sdt_xml("bid.material:business_license", "营业执照", "营业执照", 5010)))

    p = doc.add_paragraph()
    p._p.append(parse_xml(sdt_xml("bid.slot:body", "标书正文", "标书正文", 5011)))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_template_docx_role_only():
    """只含技术册插槽的模板（多册拆分验证用）。"""
    doc = Document()
    p = doc.add_paragraph()
    p._p.append(
        parse_xml(sdt_xml("bid.slot:role.technical_solution", "技术册", "技术册", 6001))
    )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def seed_data():
    """ORM 种子：项目/标段/大纲/章节/企业/材料包/材料图片。"""
    from django.contrib.auth import get_user_model

    from apps.common.services.storage import StorageService
    from apps.enterprise.models import (
        BidMaterialPackage,
        BidMaterialPackageItem,
        CompanyMaterial,
        CompanyProfile,
    )
    from apps.outline.models import Outline, Section
    from apps.projects.models import (
        Lot,
        Project,
        ProjectMember,
        ProjectRole,
    )

    User = get_user_model()
    admin = User.objects.get(username="admin")

    ts = int(time.time())
    project = Project.objects.create(name=f"E2E测试项目{ts}", created_by=admin)
    # 大纲接口按项目成员过滤，需把 admin 加为成员
    role = ProjectRole.objects.create(
        project=project, name="负责人", code=f"owner-{ts}",
        permissions=["project.view", "outline.view", "outline.edit"],
        is_builtin=True,
    )
    ProjectMember.objects.create(project=project, user=admin, project_role=role)
    lot = Lot.objects.create(name="标段一", code="01", project=project)
    outline = Outline.objects.create(
        project=project, lot=lot, name=f"E2E标书{ts}",
        source="preset", created_by=admin,
    )
    Section.objects.create(
        outline=outline, title="项目概述", level=1, sort_order=0,
        content="本项目为**重点**建设项目\n\n- 要点一\n- 要点二",
    )
    Section.objects.create(
        outline=outline, title="实施计划", level=1, sort_order=1,
        content="| 阶段 | 工期 |\n|---|---|\n| 一期 | 30天 |",
    )

    company = CompanyProfile.objects.create(
        name=f"E2E测试科技有限公司{ts}",
        unified_social_credit_code=f"91330100E2E{ts}",
        legal_representative="测试法人",
        created_by=admin,
    )
    package = BidMaterialPackage.objects.create(
        outline=outline, company=company, name="E2E材料包",
    )

    # 材料图片进真实 MinIO
    object_key = f"e2e/license-{ts}.png"
    StorageService().put_object(object_key, PNG_1X1, content_type="image/png")
    material = CompanyMaterial.objects.create(
        company=company, material_type="qualification",
        title="营业执照", object_key=object_key,
    )
    BidMaterialPackageItem.objects.create(
        package=package, material=material, usage_key="business_license",
    )
    return outline, company


def docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main():
    print("== 0. 登录 ==")
    resp = requests.post(f"{BASE}/api/auth/login", json=ADMIN, timeout=15)
    check("登录", resp.status_code == 200, resp.text[:200])
    token = resp.json()["access"]
    headers = {"Authorization": f"Bearer {token}"}

    print("== 1. 种子数据 ==")
    outline, company = seed_data()
    print(f"  outline_id={outline.id}, company={company.name}")

    print("== 2. 上传模板（含内容控件） ==")
    template_bytes = build_template_docx()
    resp = requests.post(
        f"{BASE}/api/bid-word-templates/",
        data={"name": f"E2E模板{int(time.time())}", "scope_type": "system"},
        files={"file": ("e2e.docx", template_bytes)},
        headers=headers,
        timeout=30,
    )
    check("创建模板", resp.status_code == 201, resp.text[:300])
    tpl = resp.json()
    tpl_id = tpl["id"]

    print("== 3. 模板校验（五层） ==")
    resp = requests.post(
        f"{BASE}/api/bid-word-templates/{tpl_id}/validate/", headers=headers, timeout=60,
    )
    result = resp.json()
    check("校验通过", result.get("valid") is True, str(result.get("errors")))
    check("识别变量含 company.name", "company.name" in result.get("variables", []))
    check("识别材料变量", "material:business_license" in result.get("variables", []))
    check("样式报告非空", len(result.get("styles", [])) > 0)

    print("== 4. 发布（真实 OnlyOffice 转换预览） ==")
    resp = requests.post(
        f"{BASE}/api/bid-word-templates/{tpl_id}/publish/", headers=headers, timeout=180,
    )
    check("发布成功", resp.status_code == 201, resp.text[:300])
    version = resp.json()["version"]
    check("版本号 V1", version["version_no"] == 1)
    check(
        "预览图已生成(Conversion API)",
        version.get("has_preview_image") is True,
        "has_preview_image=False，检查 ONLYOFFICE 转换链路",
    )
    check(
        "预览 PDF 已生成",
        version.get("has_preview_pdf") is True,
        "has_preview_pdf=False",
    )

    print("== 5. 选模板生成标书 ==")
    resp = requests.post(
        f"{BASE}/api/outlines/{outline.id}/build_docx/",
        json={"template_id": tpl_id},
        headers=headers,
        timeout=120,
    )
    check("生成成功", resp.status_code == 200, resp.text[:300])
    data = resp.json()
    check("返回模板信息", data.get("template", {}).get("id") == tpl_id)
    doc_id = data["document_id"]

    print("== 6. 下载 docx 并断言内容 ==")
    resp = requests.get(
        f"{BASE}/api/bid-documents/{doc_id}/download/", headers=headers, timeout=60,
    )
    check("下载成功", resp.status_code == 200)
    content = resp.content
    text = docx_text(content)
    check("项目名称变量已替换", f"E2E测试项目" in text)
    check("企业名称变量已替换", company.name in text)
    check("法定代表人变量已替换", "测试法人" in text)
    check("章节标题已插入", "项目概述" in text)
    check("正文内容已插入", "要点一" in text)
    check("表格已插入", "30天" in text)
    check("无残留渲染指令", "{{" not in text and "bid." not in text)

    import zipfile

    with zipfile.ZipFile(BytesIO(content)) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    check("材料图片已嵌入", len(media) > 0)

    print("== 7. 导出 PDF（真实转换） ==")
    resp = requests.get(
        f"{BASE}/api/bid-documents/{doc_id}/export_pdf/", headers=headers, timeout=300,
    )
    check("PDF 导出成功", resp.status_code == 200, resp.text[:200] if resp.status_code != 200 else "")
    check("PDF 魔数正确", resp.content.startswith(b"%PDF"))

    print("== 8. 多册分槽渲染 ==")
    # 给两个章节打角色，用只含技术册插槽的模板渲染
    from apps.outline.models import Section as SectionModel

    sec1 = SectionModel.objects.get(outline=outline, title="项目概述")
    sec1.content_matrix = {"section_role": "technical_solution"}
    sec1.save(update_fields=["content_matrix"])
    sec2 = SectionModel.objects.get(outline=outline, title="实施计划")
    sec2.content_matrix = {"section_role": "business_response"}
    sec2.save(update_fields=["content_matrix"])

    resp = requests.post(
        f"{BASE}/api/bid-word-templates/",
        data={"name": f"E2E分册模板{int(time.time())}", "scope_type": "system"},
        files={"file": ("ms.docx", build_template_docx_role_only())},
        headers=headers,
        timeout=30,
    )
    ms_tpl_id = resp.json()["id"]
    resp = requests.post(
        f"{BASE}/api/bid-word-templates/{ms_tpl_id}/publish/",
        headers=headers, timeout=180,
    )
    check("分册模板发布", resp.status_code == 201, resp.text[:200])
    resp = requests.post(
        f"{BASE}/api/outlines/{outline.id}/build_docx/",
        json={"template_id": ms_tpl_id},
        headers=headers,
        timeout=120,
    )
    ms_doc_id = resp.json()["document_id"]
    resp = requests.get(
        f"{BASE}/api/bid-documents/{ms_doc_id}/download/", headers=headers, timeout=60,
    )
    ms_text = docx_text(resp.content)
    check("技术册章节已渲染", "要点一" in ms_text)
    check("商务册章节被排除", "30天" not in ms_text)

    print()
    print(f"通过 {len(PASSED)} 项，失败 {len(FAILED)} 项")

    # 自清理：删除本次运行产生的数据
    print("== 清理 ==")
    try:
        from apps.outline.models import BidWordTemplate
        from apps.projects.models import Project, ProjectMember, ProjectRole

        project = outline.project
        ProjectMember.objects.filter(project=project).delete()
        ProjectRole.objects.filter(project=project).delete()
        project.delete()
        BidWordTemplate.objects.filter(id__in=[tpl_id, ms_tpl_id]).delete()
        company.delete()
        print("  已清理 E2E 数据")
    except Exception as exc:
        print(f"  清理失败（不影响测试结果）: {exc}")

    if FAILED:
        for name, detail in FAILED:
            print(f"  失败: {name} {detail}")
        sys.exit(1)


if __name__ == "__main__":
    main()
