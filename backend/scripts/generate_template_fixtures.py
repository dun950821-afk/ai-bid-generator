# backend/scripts/generate_template_fixtures.py
"""生成模板中心测试 fixtures（方案 §66）。

用法：cd backend && .venv/bin/python scripts/generate_template_fixtures.py
输出：apps/outline/tests/fixtures/templates/*.docx

fixtures 由代码生成而非手写二进制，保证可复现、可审查；
模板结构变化时重新运行本脚本即可。
"""

import os
import sys
from io import BytesIO

import django

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings.dev")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
django.setup()

from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls, qn  # noqa: E402
from docx.shared import Mm  # noqa: E402

OUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "apps/outline/tests/fixtures/templates",
)

PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def sdt(tag, alias, text, sdt_id):
    return (
        f"<w:sdt {nsdecls('w')}>"
        f'<w:sdtPr><w:alias w:val="{alias}"/><w:tag w:val="{tag}"/>'
        f'<w:id w:val="{sdt_id}"/></w:sdtPr>'
        f'<w:sdtContent><w:r><w:t>{text}</w:t></w:r></w:sdtContent>'
        f"</w:sdt>"
    )


def build(controls, plain_lines=(), header_logo=False) -> bytes:
    doc = Document()
    if header_logo:
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(BytesIO(PNG_1X1), width=Mm(20))
        inline = paragraph._p.find(".//" + qn("wp:inline"))
        inline.find(qn("wp:docPr")).set("descr", "bid.image:company.logo")
    for line in plain_lines:
        doc.add_paragraph(line)
    for i, (tag, alias, text) in enumerate(controls):
        p = doc.add_paragraph()
        p._p.append(parse_xml(sdt(tag, alias, text, 7000 + i)))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


FIXTURES = {
    # 基础模板：变量 + 正文插槽
    "simple.docx": build([
        ("bid.var:project.name", "项目名称", "项目名称"),
        ("bid.var:company.name", "企业名称", "企业名称"),
        ("bid.slot:body", "标书正文", "标书正文"),
    ]),
    # 封面 + 页眉 Logo + 正文插槽
    "cover_header_footer.docx": build(
        [
            ("bid.var:project.name", "项目名称", "项目名称"),
            ("bid.slot:body", "标书正文", "标书正文"),
        ],
        header_logo=True,
    ),
    # 多册分槽模板
    "multi_slot.docx": build([
        ("bid.slot:role.technical_solution", "技术册", "技术册"),
        ("bid.slot:role.business_response", "商务册", "商务册"),
    ]),
    # 含材料控件
    "material.docx": build([
        ("bid.material:business_license", "营业执照", "营业执照"),
        ("bid.slot:body", "标书正文", "标书正文"),
    ]),
    # 非法：无正文插槽
    "invalid_no_body.docx": build([
        ("bid.var:company.name", "企业名称", "企业名称"),
    ]),
    # 非法：未知变量
    "invalid_variable.docx": build([
        ("bid.var:company.not_exist", "不存在", "不存在"),
        ("bid.slot:body", "标书正文", "标书正文"),
    ]),
}


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    for name, content in FIXTURES.items():
        path = os.path.join(OUT_DIR, name)
        with open(path, "wb") as f:
            f.write(content)
        print(f"written: {name} ({len(content)} bytes)")


if __name__ == "__main__":
    main()
