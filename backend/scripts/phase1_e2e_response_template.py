# -*- coding: utf-8 -*-
"""Phase 1 端到端验证: 常熟农商银行文件 → 识别 → 确认 → 填充 → 产物。

用法(容器内):
    python scripts/phase1_e2e_response_template.py
"""

import os
import sys

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

import django

django.setup()

from apps.response_template.constants import TemplateStatus
from apps.response_template.models import TenderResponseDocument, TenderResponseTemplate
from apps.tender.models import TenderFile

TENDER_FILE_ID = 44  # 常熟农商银行


def main():
    tf = TenderFile.objects.get(pk=TENDER_FILE_ID)
    print(f"[e2e] 源文件: {tf.original_name}")

    # 1. 创建模板(与 API create 等价)
    pd = tf.parsed_documents.filter(is_active=True).first()
    template, _ = TenderResponseTemplate.objects.get_or_create(
        source_file=tf,
        defaults={
            "project": tf.project,
            "lot": tf.lot,
            "parsed_document": pd,
            "name": f"{tf.original_name} 响应模板",
            "status": TemplateStatus.PENDING,
        },
    )
    # 重建为待识别
    template.status = TemplateStatus.PENDING
    template.save(update_fields=["status"])

    # 2. 识别(同步执行, 等价于 celery 任务体)
    from apps.response_template.services.analyzer import ResponseTemplateAnalyzer

    summary = ResponseTemplateAnalyzer().analyze(template)
    print("\n[识别] 统计:")
    for k, v in summary.items():
        print(f"  {k}: {v}")

    blocks = template.blocks.all()
    print(f"\n[识别] 块数: {len(blocks)}")
    from collections import Counter

    dist = Counter(b.block_type for b in blocks)
    for t, c in dist.most_common():
        print(f"  {t}: {c}")

    if template.status != TemplateStatus.ANALYZED:
        print(f"\n✗ 识别失败: {template.error_message}")
        sys.exit(1)

    # 3. 确认
    template.status = TemplateStatus.CONFIRMED
    template.save(update_fields=["status"])
    print("\n[确认] 已确认, 状态:", template.status)

    # 4. 生成
    from apps.response_template.services.filler import OoxmlFiller

    content_file, warnings, filled = OoxmlFiller().fill(template, list(blocks))
    print(f"\n[生成] 文件: {content_file.name} ({content_file.size} bytes)")
    print(f"[生成] 填充成功块: {len(filled)}/{len(blocks)}")
    print(f"[生成] 警告数: {len(warnings)}")
    for w in warnings[:20]:
        print(f"  ⚠ {w['block_key']}: {w['message'][:60]}")

    # 5. 保存产物到 MinIO(等价于 fill task)
    from apps.common.services.storage import StorageService

    storage = StorageService()
    object_key = f"projects/{tf.project_id}/response/{template.id}/e2e-response.docx"
    content_file.seek(0)
    storage.put_object(object_key, content_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    doc, _ = TenderResponseDocument.objects.update_or_create(
        template=template,
        kind="main",
        defaults={
            "title": f"{template.name}",
            "status": TenderResponseDocument.STATUS_DONE,
            "object_key": object_key,
            "file_name": content_file.name,
            "file_size": content_file.size,
        },
    )
    url = storage.presigned_get_object(object_key)
    print(f"\n[产物] document#{doc.id} 下载URL: {url}")

    # 6. 校验产物可重新打开
    from docx import Document
    import io

    data = storage.get_object(object_key)
    d = Document(io.BytesIO(data))
    print(f"[校验] 产物可正常打开, 段落数: {len(d.paragraphs)}, 表格数: {len(d.tables)}")

    # 残留占位符检查
    import re

    leftover = []
    for p in d.paragraphs:
        if re.search(r"_{2,}", p.text):
            leftover.append(p.text[:50])
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if re.search(r"_{2,}", cell.text):
                    leftover.append(cell.text[:50])
    print(f"[校验] 残留下划线空位段落数: {len(leftover)}")
    for l in leftover[:10]:
        print(f"  · {l}")

    print("\n✅ 端到端验证完成")


if __name__ == "__main__":
    main()
